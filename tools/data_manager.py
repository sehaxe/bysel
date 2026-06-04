import os
import sys
import json
import base64
import shutil
import urllib.request
import typer

# Ограничиваем потоки pyarrow для предотвращения GIL-конфликтов на выходе
try:
    import pyarrow
    pyarrow.set_cpu_count(1)
    pyarrow.set_io_thread_count(1)  # <- Исправленный метод
except ImportError:
    pass

DATA_DIR = "data_train"
IMAGES_DIR = os.path.join(DATA_DIR, "images")
JSONL_PATH = os.path.join(DATA_DIR, "dataset.jsonl")
MULTIMODAL_DIR = os.path.join(DATA_DIR, "multimodal")
MULTIMODAL_MANIFEST = os.path.join(DATA_DIR, "multimodal_manifest.jsonl")

# Инициализируем приложение Typer
app = typer.Typer()

# Реестр умных пресетов на базе Generalized Chinchilla Scaling Laws (80 байт на параметр)
PRESETS = {
    "shpak": {
        "text_limit": 768000,   
        "sft_limit": 8000,      
        "vision_limit": 1000    
    },
    "chyzh": {
        "text_limit": 130000,   
        "sft_limit": 2000,      
        "vision_limit": 200     
    }
}


def ensure_directories():
    os.makedirs(IMAGES_DIR, exist_ok=True)


def _download_vision(limit: int, dataset_name: str):
    from datasets import load_dataset
    import warnings
    warnings.filterwarnings("ignore") # Подавление фонового мусора HTTP-соединений
    
    ensure_directories()
    
    if os.path.exists(JSONL_PATH) and os.path.getsize(JSONL_PATH) > 1024:
        images_count = len(os.listdir(IMAGES_DIR)) if os.path.exists(IMAGES_DIR) else 0
        if images_count > 0:
            file_size_mb = os.path.getsize(JSONL_PATH) / (1024 * 1024)
            typer.echo(typer.style(f"📁 Vision dataset '{JSONL_PATH}' already exists ({file_size_mb:.2f} MB) with {images_count} images. Skipping download.", fg=typer.colors.GREEN))
            return

    # 🎯 ИСПРАВЛЕНИЕ COCO: Переход на идеальный Parquet-вариант с нативными PIL-изображениями
    if dataset_name in ["HuggingFaceM4/COCO", "DavidPhilips/coco2017"]:
        dataset_name = "jxie/coco_captions"

    typer.echo(typer.style(f"📥 Connecting to HF and streaming '{dataset_name}'...", fg=typer.colors.CYAN))
    try:
        dataset = load_dataset(dataset_name, split="train", streaming=True)
    except Exception as e:
        typer.echo(typer.style(f"❌ Failed to load dataset: {e}", fg=typer.colors.RED, bold=True))
        return
        
    count = 0
    with open(JSONL_PATH, "a", encoding="utf-8") as f:
        for item in dataset:
            if count >= limit:
                break
            try:
                img = item["image"]
                
                # Гибкий поиск текста подписи
                caption = ""
                if "text" in item:
                    caption = item["text"]
                elif "caption" in item:
                    caption = item["caption"]
                elif "captions" in item:
                    caption = item["captions"][0] if isinstance(item["captions"], list) else item["captions"]
                    
                caption = str(caption).strip()
                if not caption:
                    continue
                    
                img_filename = f"images/coco_{count}.jpg"
                img_path = os.path.join(DATA_DIR, img_filename)
                
                img.save(img_path)
                
                line = {"image": img_filename, "text": caption}
                f.write(json.dumps(line, ensure_ascii=False) + "\n")
                count += 1
                if count % 50 == 0:
                    typer.echo(f"   Downloaded: {count}/{limit} images...")
            except Exception:
                continue
                
    typer.echo(typer.style(f"✅ Successfully saved {count} samples to '{JSONL_PATH}'", fg=typer.colors.GREEN))

    # Избегаем GIL-конфликтов при очистке потоков PyArrow
    del dataset
    import gc
    gc.collect()


def _download_text(limit: int, source: str):
    from datasets import load_dataset
    import warnings
    warnings.filterwarnings("ignore")
    
    ensure_directories()
    source_clean = source.lower().strip()
    
    if source_clean == "tinystories":
        dataset_name, split_name, name_param, text_key = "roneneldan/TinyStories", "train", None, "text"
        output_file = os.path.join(DATA_DIR, "pretrain_tinystories.txt")
    elif source_clean == "fineweb":
        dataset_name, split_name, name_param, text_key = "HuggingFaceFW/fineweb-edu", "train", "sample-10BT", "text"
        output_file = os.path.join(DATA_DIR, "pretrain_fineweb.txt")
    elif source_clean in ["smollm", "cosmopedia"]:
        dataset_name, split_name, name_param, text_key = "HuggingFaceTB/smollm-corpus", "train", "cosmopedia-v2", "text"
        output_file = os.path.join(DATA_DIR, "pretrain_cosmopedia.txt")
    else:
        typer.echo(typer.style("❌ Unsupported source!", fg=typer.colors.RED))
        return

    if os.path.exists(output_file) and os.path.getsize(output_file) > 1024:
        file_size_mb = os.path.getsize(output_file) / (1024 * 1024)
        typer.echo(typer.style(f"📁 Pretrain file '{output_file}' already exists ({file_size_mb:.2f} MB). Skipping download.", fg=typer.colors.GREEN))
        return

    typer.echo(typer.style(f"📥 Streaming '{dataset_name}' from Hugging Face...", fg=typer.colors.CYAN))
    try:
        if name_param:
            dataset = load_dataset(dataset_name, name=name_param, split=split_name, streaming=True)
        else:
            dataset = load_dataset(dataset_name, split=split_name, streaming=True)
    except Exception as e:
        typer.echo(typer.style(f"❌ Load error: {e}", fg=typer.colors.RED))
        return

    count = 0
    with open(output_file, "a", encoding="utf-8") as f:
        for item in dataset:
            if count >= limit:
                break
            try:
                text_content = item[text_key].strip()
                if not text_content: 
                    continue
                f.write(text_content + "\n\n")
                count += 1
                if count % 2000 == 0:
                    typer.echo(f"   Saved: {count}/{limit} texts...")
            except Exception:
                continue
    typer.echo(typer.style(f"✅ Successfully saved {count} texts to '{output_file}'", fg=typer.colors.GREEN))

    # Избегаем GIL-конфликтов при очистке потоков PyArrow
    del dataset
    import gc
    gc.collect()


def _download_sft(limit: int, source: str):
    from datasets import load_dataset
    import warnings
    warnings.filterwarnings("ignore")
    
    ensure_directories()
    source_clean = source.lower().strip()
    
    if source_clean == "alpaca":
        dataset_name = "tatsu-lab/alpaca"
        output_file = os.path.join(DATA_DIR, "sft_alpaca.jsonl")
    elif source_clean == "smoltalk":
        dataset_name = "HuggingFaceTB/smoltalk"
        output_file = os.path.join(DATA_DIR, "sft_smoltalk.jsonl")
    else:
        typer.echo(typer.style("❌ Unsupported SFT source!", fg=typer.colors.RED))
        return
    
    if os.path.exists(output_file) and os.path.getsize(output_file) > 1024:
        file_size_mb = os.path.getsize(output_file) / (1024 * 1024)
        typer.echo(typer.style(f"📁 SFT file '{output_file}' already exists ({file_size_mb:.2f} MB). Skipping download.", fg=typer.colors.GREEN))
        return

    typer.echo(typer.style(f"📥 Streaming English instruction dataset '{dataset_name}'...", fg=typer.colors.CYAN))
    try:
        if source_clean == "smoltalk":
            dataset = load_dataset(dataset_name, "all", split="train", streaming=True)
        else:
            dataset = load_dataset(dataset_name, split="train", streaming=True)
    except Exception as e:
        typer.echo(typer.style(f"❌ Failed to load SFT dataset: {e}", fg=typer.colors.RED))
        return

    count = 0
    with open(output_file, "a", encoding="utf-8") as f:
        for item in dataset:
            if count >= limit:
                break
            try:
                if source_clean == "alpaca":
                    instruction = item.get("instruction", "").strip()
                    inp = item.get("input", "").strip()
                    output = item.get("output", "").strip()
                    if not instruction or not output: 
                        continue
                    full_prompt = f"User: {instruction}"
                    if inp: 
                        full_prompt += f"\nContext: {inp}"
                    full_prompt += f"\nAssistant: {output}"
                else:  
                    messages = item.get("messages", [])
                    if not messages: 
                        continue
                    full_prompt = ""
                    for msg in messages:
                        role = msg.get("role", "user").capitalize()
                        content = msg.get("content", "").strip()
                        full_prompt += f"{role}: {content}\n"
                
                line = {"text": full_prompt.strip()}
                f.write(json.dumps(line, ensure_ascii=False) + "\n")
                count += 1
                if count % 500 == 0:
                    typer.echo(f"   Converted: {count}/{limit} instructions...")
            except Exception:
                continue
    typer.echo(typer.style(f"✅ SFT dataset ({count} instructions) successfully saved to '{output_file}'", fg=typer.colors.GREEN))

    # Избегаем GIL-конфликтов при очистке потоков PyArrow
    del dataset
    import gc
    gc.collect()


# Регистрируем команды в Typer
@app.command()
def download_all(
    text_limit: int = typer.Option(5000, "--text-limit", "-t", help="Limit for pretrain text"),
    sft_limit: int = typer.Option(3000, "--sft-limit", "-s", help="Limit for SFT instructions"),
    vision_limit: int = typer.Option(1000, "--vision-limit", "-v", help="Limit for COCO images"),
    preset: str = typer.Option(None, "--preset", "-p", help="Automatic profile preset: 'shpak'")
):
    if preset:
        preset_clean = preset.lower().strip()
        if preset_clean in PRESETS:
            text_limit = PRESETS[preset_clean]["text_limit"]
            sft_limit = PRESETS[preset_clean]["sft_limit"]
            vision_limit = PRESETS[preset_clean]["vision_limit"]
            typer.echo(typer.style(f"🦁 PRESET DETECTED: {preset_clean}", fg=typer.colors.GREEN, bold=True))
        else:
            typer.echo(typer.style(f"⚠️ Unknown preset '{preset}'! Using manual limits.", fg=typer.colors.YELLOW))

    typer.echo(typer.style("\n📥 STARTING BULK DATASET DOWNLOAD...", fg=typer.colors.CYAN, bold=True))
    _download_text(text_limit, "smollm")
    _download_sft(sft_limit, "smoltalk")
    _download_vision(vision_limit, "HuggingFaceM4/COCO")


@app.command()
def download_vision(
    limit: int = typer.Option(1000, "--limit", "-l", help="Number of images to download"),
    dataset_name: str = typer.Option("HuggingFaceM4/COCO", "--dataset", "-d", help="Hugging Face dataset name")
):
    _download_vision(limit, dataset_name)


@app.command()
def download_text(
    limit: int = typer.Option(5000, "--limit", "-l", help="Number of pretrain texts to download"),
    source: str = typer.Option("smollm", "--source", "-s", help="Source: 'smollm', 'fineweb', 'tinystories'"),
    preset: str = typer.Option(None, "--preset", "-p", help="Automatic profile preset: 'shpak'")
):
    if preset:
        preset_clean = preset.lower().strip()
        if preset_clean in PRESETS:
            limit = PRESETS[preset_clean]["text_limit"]
    _download_text(limit, source)


@app.command()
def download_sft(
    limit: int = typer.Option(3000, "--limit", "-l", help="Number of SFT instructions to download"),
    source: str = typer.Option("smoltalk", "--source", "-s", help="SFT Source: 'smoltalk' or 'alpaca'"),
    preset: str = typer.Option(None, "--preset", "-p", help="Automatic profile preset: 'shpak'")
):
    if preset:
        preset_clean = preset.lower().strip()
        if preset_clean in PRESETS:
            limit = PRESETS[preset_clean]["sft_limit"]
    _download_sft(limit, source)


@app.command()
def label_vision(
    source_dir: str = typer.Option("my_photos", "--dir", "-s", help="Directory with raw images"),
    model: str = typer.Option("moondream", "--model", "-m", help="Local vision model in Ollama")
):
    pass


def _synth_image(path: str, w: int, h: int, seed: int):
    import cv2
    import numpy as np
    rng = np.random.default_rng(seed)
    arr = rng.integers(0, 255, (h, w, 3), dtype=np.uint8)
    arr = cv2.putText(arr, f"busel #{seed}", (4, 16), cv2.FONT_HERSHEY_SIMPLEX, 0.4, (255, 255, 255), 1, cv2.LINE_AA)
    cv2.imwrite(path, arr)


def _synth_video(path: str, n_frames: int, w: int, h: int, seed: int):
    import cv2
    import numpy as np
    rng = np.random.default_rng(seed)
    fourcc = cv2.VideoWriter_fourcc(*"mp4v")
    writer = cv2.VideoWriter(path, fourcc, 10.0, (w, h))
    for i in range(n_frames):
        frame = rng.integers(0, 255, (h, w, 3), dtype=np.uint8)
        frame = cv2.putText(frame, f"f{i}", (4, 16), cv2.FONT_HERSHEY_SIMPLEX, 0.4, (255, 255, 255), 1, cv2.LINE_AA)
        writer.write(frame)
    writer.release()


def _synth_audio(path: str, seconds: float, sr: int, seed: int):
    import numpy as np
    import soundfile as sf
    rng = np.random.default_rng(seed)
    n = int(seconds * sr)
    data = (rng.standard_normal(n) * 0.1).astype(np.float32)
    sf.write(path, data, sr)


def _synth_docx(path: str, seed: int):
    import docx
    d = docx.Document()
    d.add_heading(f"busel multimodal test #{seed}", level=1)
    d.add_paragraph("This is a synthetic document for testing the docx encoder.")
    d.add_paragraph(f"Seed: {seed}; created by `download-multimodal`.")
    d.save(path)


def _download_multimodal(n_per_kind: int):
    os.makedirs(MULTIMODAL_DIR, exist_ok=True)

    has_docx = True
    has_sf = True
    try:
        import docx as _dx
    except ImportError:
        typer.echo(typer.style("⚠️ python-docx missing — docx samples skipped", fg=typer.colors.YELLOW))
        has_docx = False
    try:
        import soundfile as _sf
    except ImportError:
        typer.echo(typer.style("⚠️ soundfile missing — audio samples skipped", fg=typer.colors.YELLOW))
        has_sf = False

    manifest = []
    typer.echo(typer.style(f"🛰️ Generating {n_per_kind} synthetic samples per modality in '{MULTIMODAL_DIR}/'...", fg=typer.colors.CYAN, bold=True))

    for i in range(n_per_kind):
        img_path = os.path.join(MULTIMODAL_DIR, f"img_{i}.png")
        _synth_image(img_path, 64, 64, seed=i)
        manifest.append({"path": os.path.relpath(img_path, DATA_DIR), "modality": "image", "caption": f"synthetic image #{i}"})
        if i % max(1, n_per_kind // 4) == 0:
            typer.echo(f"   image: {i + 1}/{n_per_kind}")

    for i in range(n_per_kind):
        vid_path = os.path.join(MULTIMODAL_DIR, f"vid_{i}.mp4")
        _synth_video(vid_path, n_frames=12, w=64, h=64, seed=i + 1000)
        manifest.append({"path": os.path.relpath(vid_path, DATA_DIR), "modality": "video", "caption": f"synthetic video #{i}"})
        if i % max(1, n_per_kind // 4) == 0:
            typer.echo(f"   video: {i + 1}/{n_per_kind}")

    if has_sf:
        for i in range(n_per_kind):
            wav_path = os.path.join(MULTIMODAL_DIR, f"aud_{i}.wav")
            _synth_audio(wav_path, seconds=1.0, sr=16000, seed=i + 2000)
            manifest.append({"path": os.path.relpath(wav_path, DATA_DIR), "modality": "audio", "caption": f"synthetic audio #{i}"})
            if i % max(1, n_per_kind // 4) == 0:
                typer.echo(f"   audio: {i + 1}/{n_per_kind}")

    if has_docx:
        for i in range(n_per_kind):
            docx_path = os.path.join(MULTIMODAL_DIR, f"doc_{i}.docx")
            _synth_docx(docx_path, seed=i + 3000)
            manifest.append({"path": os.path.relpath(docx_path, DATA_DIR), "modality": "docx", "caption": f"synthetic docx #{i}"})
            if i % max(1, n_per_kind // 4) == 0:
                typer.echo(f"   docx: {i + 1}/{n_per_kind}")

    typer.echo(typer.style("⚠️ PDF: requires `uv add docling` (heavyweight). Install separately, then drop into multimodal/ to train on PDFs.", fg=typer.colors.YELLOW))

    with open(MULTIMODAL_MANIFEST, "w", encoding="utf-8") as f:
        for entry in manifest:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    typer.echo(typer.style(f"✅ Wrote {len(manifest)} entries to '{MULTIMODAL_MANIFEST}'", fg=typer.colors.GREEN, bold=True))
    typer.echo(typer.style(f"📂 Files: {MULTIMODAL_DIR}/", fg=typer.colors.GREEN))


@app.command(name="download-multimodal")
def download_multimodal(
    limit: int = typer.Option(8, "--limit", "-l", help="Number of synthetic samples per modality (image, video, audio, docx)")
):
    """🛰️ Generate synthetic image/video/audio/docx test files for the multimodal encoders (no internet needed)."""
    _download_multimodal(limit)


# Безопасный запуск программы с принудительной выгрузкой зависших фоновых потоков
if __name__ == "__main__":
    try:
        app()
    finally:
        # Сброс буферов вывода
        sys.stdout.flush()
        sys.stderr.flush()
        # Принудительное закрытие процесса на уровне ОС во избежание зависания
        os._exit(0)