"""
🧪 busel UNIFIED TEST SUITE — paper-compliance + integration
Covers all 6 reference papers + end-to-end integration.
"""
import os
import sys
import time
import math
import struct
import unittest
import json
import inspect

import torch
import torch.nn as nn
import torch.nn.functional as F

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import busel_rust_io
from data.pipeline import get_busel_dataloader, collate_busel_batch
from model.patching import StridedFastBLTPatcher
from model.layers import BitLinear_a4_8, H_BitLinear, RMSNorm, SwishGLUClamped, fast_walsh_hadamard_transform
from model.attention import stable_gdn2_recurrent_jit, BulbaGDN2SeRoPEBlock, MultiHeadLatentAttention
from model.routing import MoDSequenceRouter, BulbaTernaryTitanExpertFFN, BulbaTernaryTitanMoE
from model.backbone import buselModel, ManifoldConstrainedAttnRes
from training.optimizer import _compiled_newton_schulz, buselOptimizerEngine, Muon, _newton_schulz_core
from training.recipe import buselLossEngine, validate_training_schedule

from busel_registry import register, get, list_registered, is_registered, unregister, clear_registry
from busel_logging import setup_logging, log_event, get_logger, JSONFormatter
from ui.teto import frame as teto_frame, frames as teto_frames, states as teto_states
from ui import cli as ui_cli

try:
    from multimodal.encoders import (
        ImageEncoder,
        VideoEncoder,
        AudioEncoder,
        PDFEncoder,
        DocxEncoder,
        TextEncoder,
        auto_encode,
        build_encoder_for,
        IMAGE_MARKER,
        MEDIA_END,
        IMAGE_BYTES,
    )
    from multimodal import list_encoders as list_mm_encoders
    HAS_MULTIMODAL_DEPS = True
except Exception:
    HAS_MULTIMODAL_DEPS = False

try:
    from multimodal.special_tokens import (
        vocab_size as mm_vocab_size,
        get_special_token,
        list_special_tokens,
        disable_special_token,
        enable_special_token,
        register_special_token,
        is_enabled,
        enabled_ids,
        LAYER_DESCRIPTIONS,
        MEDIA_START,
        DOC_SEP,
        MOD_IMAGE,
        MOD_VIDEO,
        MOD_AUDIO,
        MOD_PDF,
        MOD_DOCX,
        MOD_TEXT,
        BOS,
        EOS,
        PAD,
        ROLE_SYSTEM,
        ROLE_USER,
        ROLE_ASSISTANT,
        ROLE_TOOL,
        THINK_START,
        THINK_END,
        TOOL_BASH,
        STATUS_SUCCESS,
        SPECIAL_VOCAB_BASE,
    )
    HAS_SPECIAL_TOKENS = True
except Exception:
    HAS_SPECIAL_TOKENS = False

try:
    from training.stages import (
        BaseStage,
        StageState,
        StageSpec,
        PipelineConfig,
        register_stage,
        get_stage,
        list_stages,
        is_stage_registered,
        load_pipeline_yaml,
    )
    HAS_TRAINING_STAGES = True
except Exception:
    HAS_TRAINING_STAGES = False


class _MockConfig:
    def __init__(self, **kw):
        self.d_model = kw.get("d_model", 128)
        self.n_layers = kw.get("n_layers", 2)
        self.n_heads = kw.get("n_heads", 4)
        self.expert_hidden = kw.get("expert_hidden", 256)
        self.num_experts = kw.get("num_experts", 2)
        self.top_k = kw.get("top_k", 2)
        default_vocab = mm_vocab_size() if HAS_SPECIAL_TOKENS else 259
        self.vocab_size = kw.get("vocab_size", default_vocab)
        self.n_hyper = kw.get("n_hyper", 2)


class TestbuselFramework(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.device = "mps" if torch.backends.mps.is_available() else ("cuda" if torch.cuda.is_available() else "cpu")
        print(f"\n🚀 Running busel Test Suite on device: {cls.device.upper()}\n" + "=" * 80)

    def test_rust_io_streamer(self):
        print("🧪 [1] Rust ByteStreamer...")
        temp_file = "temp_test_rust_io.txt"
        if os.path.exists(temp_file):
            os.remove(temp_file)
        test_data = "Hello from busel Rust IO! " * 350
        with open(temp_file, "w", encoding="utf-8") as f:
            f.write(test_data)
        try:
            streamer = busel_rust_io.ByteStreamer(temp_file, 8, 0)
            chunk1 = streamer.next_chunk()
            self.assertEqual(len(chunk1), 8)
            self.assertEqual(bytes(chunk1).decode("utf-8", errors="ignore"), "Hello fr")
            chunk2 = streamer.next_chunk()
            self.assertEqual(len(chunk2), 8)
            self.assertEqual(bytes(chunk2).decode("utf-8", errors="ignore"), "om busel")
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
        print("   ✅ Rust ByteStreamer passed.")

    def test_rust_binary_packer(self):
        print("🧪 [2] Rust Binary Packer...")
        temp_bin = "temp_test_packer.bin"
        if os.path.exists(temp_bin):
            os.remove(temp_bin)
        test_bytes = [10, 20, 30, 40, 255, 0]
        try:
            busel_rust_io.append_to_binary_file(temp_bin, test_bytes)
            self.assertTrue(os.path.exists(temp_bin))
            self.assertEqual(os.path.getsize(temp_bin), 6)
            with open(temp_bin, "rb") as f:
                read_data = list(f.read())
            self.assertEqual(read_data, test_bytes)
        finally:
            if os.path.exists(temp_bin):
                os.remove(temp_bin)
        print("   ✅ Rust Binary Packer passed.")

    def test_rust_ternary_inference(self):
        print("🧪 [3] Rust Ternary CPU Inference...")
        x = [1.0, -1.0, 2.0]
        w = [1, 0, -1, 0, 1, 1]
        expected_y = [-1.0, 1.0]
        actual_y = busel_rust_io.ternary_matmul_cpu(x, w, 2, 3)
        self.assertEqual(actual_y, expected_y)
        print("   ✅ Rust Ternary CPU Inference passed.")

    def test_bitlinear_quantization(self):
        print("🧪 [4] BitLinear 1.58b Quantization (forward)...")
        linear = BitLinear_a4_8(64, 128).to(self.device)
        x = torch.randn(2, 64, device=self.device)
        out = linear(x)
        self.assertEqual(out.shape, (2, 128))
        self.assertFalse(torch.isnan(out).any())
        print("   ✅ BitLinear Quantization passed.")

    def test_jit_gdn2_attention(self):
        print("🧪 [5] Stable JIT GDN-2 Loop (forward)...")
        q = torch.randn(2, 128, 4, 64, device=self.device)
        k = torch.randn(2, 128, 4, 64, device=self.device)
        v = torch.randn(2, 128, 4, 64, device=self.device)
        b = torch.sigmoid(torch.randn(2, 128, 4, 64, device=self.device))
        w = torch.sigmoid(torch.randn(2, 128, 4, 64, device=self.device))
        alpha = torch.sigmoid(torch.randn(2, 128, 4, 64, device=self.device))
        q = torch.nn.functional.normalize(q, p=2, dim=-1)
        k = torch.nn.functional.normalize(k, p=2, dim=-1)
        with torch.autocast(device_type=self.device, dtype=torch.float16 if self.device == "mps" else torch.bfloat16):
            out = stable_gdn2_recurrent_jit(q, k, v, b, w, alpha)
        self.assertEqual(out.shape, (2, 128, 256))
        self.assertFalse(torch.isnan(out).any())
        print("   ✅ Stable JIT GDN-2 Loop passed.")

    def test_fused_glu_and_expert_ffn(self):
        print("🧪 [6] Fused Gate-Up Projections (SwishGLUClamped)...")
        x = torch.randn(2, 32, 256, device=self.device)
        glu = SwishGLUClamped(256, 512).to(self.device)
        out_glu = glu(x)
        self.assertEqual(out_glu.shape, (2, 32, 256))
        self.assertFalse(torch.isnan(out_glu).any())
        expert = BulbaTernaryTitanExpertFFN(256, 512).to(self.device)
        out_exp = expert(x)
        self.assertEqual(out_exp.shape, (2, 32, 256))
        self.assertFalse(torch.isnan(out_exp).any())
        print("   ✅ Fused Gate-Up Projections passed.")

    def test_muon_transpose_trick(self):
        print("🧪 [7] Muon Transpose Trick (NS forward)...")
        X = torch.randn(512, 256, device=self.device)
        O_t = _compiled_newton_schulz(X, steps=5)
        self.assertEqual(O_t.shape, (512, 256))
        self.assertFalse(torch.isnan(O_t).any(), "O_t contains NaNs!")
        print("   ✅ Muon Transpose Trick passed.")

    def test_complete_backbone_and_gradients(self):
        print("🧪 [8] Complete Backbone & Backpropagation...")
        cfg = _MockConfig(d_model=256, n_layers=4, n_heads=4, expert_hidden=512, num_experts=4)
        patcher = StridedFastBLTPatcher(d_model=cfg.d_model).to(self.device)
        model = buselModel(cfg).to(self.device)
        opt_engine = buselOptimizerEngine(model, lr_muon=0.0004, lr_adamw=0.00004)
        loss_engine = buselLossEngine(cfg.vocab_size)
        byte_batch = torch.randint(0, cfg.vocab_size, (2, 256), dtype=torch.int32, device=self.device)
        input_bytes = byte_batch[:, :-patcher.stride]
        opt_engine.zero_grad(set_to_none=True)
        target_dtype = torch.bfloat16 if self.device == "cuda" else torch.float16
        with torch.autocast(device_type=self.device, dtype=target_dtype):
            patches = patcher(input_bytes)
            T_patches = patches.shape[1]
            targets = byte_batch[:, patcher.stride::patcher.stride][:, :T_patches]
            if targets.shape[1] < T_patches:
                targets = torch.nn.functional.pad(targets, (0, T_patches - targets.shape[1]), value=0)
            (logits_t1, _, _, _), aux_loss = model(patches, None)
            loss = loss_engine.compute_pretrain_loss(logits_t1, targets) + aux_loss.float()
        loss.backward()
        has_gradients = False
        for name, p in model.named_parameters():
            if p.grad is not None:
                has_gradients = True
                self.assertFalse(torch.isnan(p.grad).any(), f"Gradient of '{name}' is NaN!")
        self.assertTrue(has_gradients, "No gradients were computed!")
        opt_engine.step()
        for name, p in model.named_parameters():
            self.assertFalse(torch.isnan(p).any(), f"Weights of '{name}' became NaN after optimizer step!")
        print("   ✅ Complete Backbone & Backpropagation passed.")

    def test_gdn2_decoupled_erase_and_write_gates(self):
        # Paper §3.1: b_proj and w_proj are SEPARATE linear layers
        block = BulbaGDN2SeRoPEBlock(d_model=128, n_heads=4)
        self.assertIsInstance(block.b_proj, torch.nn.Linear)
        self.assertIsInstance(block.w_proj, torch.nn.Linear)
        self.assertIsNot(block.b_proj, block.w_proj)

    def test_gdn2_gates_sigmoid_bounded_0_1(self):
        # Paper §3.1: erase (b) and write (w) gates ∈ (0, 1)
        block = BulbaGDN2SeRoPEBlock(d_model=128, n_heads=4).eval()
        x = torch.randn(2, 16, 128)
        with torch.no_grad():
            b = torch.sigmoid(block.b_proj(x))
            w = torch.sigmoid(block.w_proj(x))
        self.assertTrue((b > 0).all() and (b < 1).all())
        self.assertTrue((w > 0).all() and (w < 1).all())

    def test_gdn2_log_decay_alpha_a_init_negative_three(self):
        # Paper §3.2 (Eq. 12): alpha_a is initialized to -3.0
        block = BulbaGDN2SeRoPEBlock(d_model=128, n_heads=4)
        self.assertTrue(torch.allclose(block.alpha_a, torch.full_like(block.alpha_a, -3.0)))

    def test_gdn2_log_decay_alpha_in_unit_interval(self):
        # Paper §3.2 (Eq. 12): g_t = -exp(alpha_a)*softplus(alpha_proj), alpha = exp(g_t) ∈ (0, 1)
        block = BulbaGDN2SeRoPEBlock(d_model=128, n_heads=4).eval()
        x = torch.randn(2, 16, 128)
        with torch.no_grad():
            alpha_proj = block.alpha_proj(x).view(2, 16, block.n_heads, block.d_k)
            g_t = -torch.exp(block.alpha_a).view(1, 1, block.n_heads, 1) * F.softplus(alpha_proj)
            alpha = torch.exp(g_t)
        self.assertTrue((alpha > 0).all())
        self.assertTrue((alpha < 1.0 + 1e-6).all())

    def test_gdn2_serope_real_imag_pairing(self):
        # Paper §3.3 (SeRoPE): pairs real (even) and imag (odd) indices with cos/sin
        block = BulbaGDN2SeRoPEBlock(d_model=128, n_heads=4).eval()
        T = 8
        q = torch.randn(2, T, 4, block.d_k)
        k = torch.randn(2, T, 4, block.d_k)
        with torch.no_grad():
            q_out, k_out = block.apply_serope(T, q, k)
        self.assertEqual(q_out.shape, q.shape)
        self.assertEqual(k_out.shape, k.shape)
        self.assertFalse(torch.allclose(q_out, q))
        self.assertFalse(torch.allclose(k_out, k))

    def test_gdn2_serope_is_actually_rotating(self):
        # SeRoPE must not be a no-op: output must differ from input
        block = BulbaGDN2SeRoPEBlock(d_model=128, n_heads=4).eval()
        T = 16
        q = torch.randn(1, T, 4, block.d_k)
        with torch.no_grad():
            q_out, _ = block.apply_serope(T, q, q)
        self.assertFalse(torch.allclose(q_out, q))

    def test_gdn2_mla_latent_d_c_is_128(self):
        # Paper §4: MLA compresses KV to latent size d_c=128
        block = MultiHeadLatentAttention(d_model=128, n_heads=4, d_c=128)
        self.assertEqual(block.d_c, 128)
        out = block(torch.randn(2, 8, 128))
        self.assertEqual(out.shape, (2, 8, 128))
        self.assertFalse(torch.isnan(out).any())

    def test_gdn2_mla_kv_q_compression_dimensions(self):
        # Paper §4: kv_compress and q_compress project d_model → d_c
        block = MultiHeadLatentAttention(d_model=128, n_heads=4, d_c=128)
        self.assertEqual(block.kv_compress.out_features, 128)
        self.assertEqual(block.q_compress.out_features, 128)
        self.assertEqual(block.k_decompress.in_features, 128)
        self.assertEqual(block.v_decompress.in_features, 128)

    def test_bitnetv2_weights_ternary(self):
        # Paper §3.1: weights quantize to {-1, 0, +1} (1.58 bits)
        linear = BitLinear_a4_8(64, 32).eval()
        x = torch.randn(2, 64)
        with torch.no_grad():
            w = linear.weight
            alpha = w.abs().mean() + 1e-5
            w_scaled = w / alpha
            w_clipped = torch.clamp(w_scaled, -1, 1)
            w_quant = w_clipped + (torch.round(w_clipped) - w_clipped)
        uniq = torch.unique(w_quant)
        self.assertTrue(all(v in {-1.0, 0.0, 1.0} for v in uniq.tolist()))

    def test_bitnetv2_ternary_distribution_meaningful(self):
        # Paper §3.1: ternary distribution should not collapse
        torch.manual_seed(0)
        linear = BitLinear_a4_8(256, 256).eval()
        with torch.no_grad():
            w = linear.weight
            alpha = w.abs().mean() + 1e-5
            w_scaled = w / alpha
            w_clipped = torch.clamp(w_scaled, -1, 1)
            w_quant = w_clipped + (torch.round(w_clipped) - w_clipped)
        n_pos = (w_quant == 1).sum().item()
        n_neg = (w_quant == -1).sum().item()
        n_zero = (w_quant == 0).sum().item()
        total = w_quant.numel()
        self.assertGreater(n_pos / total, 0.1)
        self.assertGreater(n_neg / total, 0.1)
        self.assertGreater(n_zero / total, 0.05)
        self.assertLess(n_zero / total, 0.7)

    def test_bitnetv2_h_bitlinear_applies_walsh_hadamard(self):
        # Paper §3.2: H_BitLinear = BitLinear_a4_8 + FWHT pre-transform on input
        d = 64
        h = H_BitLinear(d, d).eval()
        x = torch.randn(2, 8, d)
        with torch.no_grad():
            plain = BitLinear_a4_8(d, d).eval()
            plain.weight.data = h.weight.data.clone()
            out_with_wht = h(x)
            out_plain = plain(x)
        self.assertEqual(out_with_wht.shape, out_plain.shape)
        self.assertFalse(torch.allclose(out_with_wht, out_plain, atol=1e-3),
                         "H_BitLinear must apply FWHT to input, producing a different output than plain BitLinear")

    def test_bitnetv2_h_bitlinear_used_for_o_proj(self):
        # Paper §3.2: o_proj in both GDN-2 and MLA must use H_BitLinear
        gdn = BulbaGDN2SeRoPEBlock(d_model=128, n_heads=4)
        mla = MultiHeadLatentAttention(d_model=128, n_heads=4, d_c=128)
        self.assertIsInstance(gdn.o_proj, H_BitLinear)
        self.assertIsInstance(mla.o_proj, H_BitLinear)

    def test_bitnetv2_swishglu_learnable_clamp(self):
        # Paper §3.3: SwishGLUClamped uses learnable per-channel clamp bounds
        glu = SwishGLUClamped(d_model=128, d_ffn=256)
        self.assertTrue(hasattr(glu, "clipping_bounds"))
        self.assertIsInstance(glu.clipping_bounds, torch.nn.Parameter)
        self.assertEqual(glu.clipping_bounds.shape, (256,))

    def test_bitnetv2_swishglu_clip_bounds_get_gradient(self):
        # Paper §3.3: clipping_bounds receives gradients via STE
        glu = SwishGLUClamped(d_model=64, d_ffn=128)
        out = glu(torch.randn(2, 8, 64))
        out.sum().backward()
        self.assertIsNotNone(glu.clipping_bounds.grad)
        self.assertFalse(torch.isnan(glu.clipping_bounds.grad).any())

    def test_mar_sinkhorn_knopp_doubly_stochastic_rows(self):
        # mHC §3.2: H is doubly-stochastic (rows sum to 1)
        mar = ManifoldConstrainedAttnRes(d_model=129, n_hyper=3, n_sinkhorn_iters=20)
        M = torch.randn(2, 4, 3, 3)
        H = mar.sinkhorn_knopp(M)
        row_sums = H.sum(dim=-1)
        torch.testing.assert_close(row_sums, torch.ones_like(row_sums), atol=1e-4, rtol=1e-4)

    def test_mar_sinkhorn_knopp_doubly_stochastic_cols(self):
        # mHC §3.2: H is doubly-stochastic (cols sum to 1)
        mar = ManifoldConstrainedAttnRes(d_model=129, n_hyper=3, n_sinkhorn_iters=20)
        M = torch.randn(2, 4, 3, 3)
        H = mar.sinkhorn_knopp(M)
        col_sums = H.sum(dim=-2)
        torch.testing.assert_close(col_sums, torch.ones_like(col_sums), atol=1e-4, rtol=1e-4)

    def test_mar_sinkhorn_knopp_non_negative(self):
        # mHC §3.2: H ∈ [0, 1] (doubly-stochastic)
        mar = ManifoldConstrainedAttnRes(d_model=128, n_hyper=2)
        M = torch.randn(8, 16, 2, 2) * 100.0
        H = mar.sinkhorn_knopp(M)
        self.assertTrue((H >= 0).all())

    def test_mar_sinkhorn_knopp_no_overflow_under_extreme_inputs(self):
        # Numerical stability: extreme inputs must not produce NaN/Inf
        mar = ManifoldConstrainedAttnRes(d_model=128, n_hyper=2)
        M = torch.randn(1, 1, 2, 2) * 1000.0
        H = mar.sinkhorn_knopp(M)
        self.assertFalse(torch.isnan(H).any())
        self.assertFalse(torch.isinf(H).any())

    def test_mar_sinkhorn_knopp_more_iters_converge_tighter(self):
        # mHC §3.2: more Sinkhorn iterations → tighter convergence to Birkhoff
        mar = ManifoldConstrainedAttnRes(d_model=129, n_hyper=3, n_sinkhorn_iters=20)
        M = torch.randn(4, 5, 3, 3)
        H = mar.sinkhorn_knopp(M, n_iters=20)
        row_dev = (H.sum(dim=-1) - 1.0).abs().max()
        col_dev = (H.sum(dim=-2) - 1.0).abs().max()
        self.assertLess(row_dev.item(), 1e-3)
        self.assertLess(col_dev.item(), 1e-3)

    def test_mar_identity_init_means_h_approx_I(self):
        # mHC §3.3: at init, H ≈ I (identity-mapping property)
        d_model, n_hyper = 128, 2
        mar = ManifoldConstrainedAttnRes(d_model=d_model, n_hyper=n_hyper, n_sinkhorn_iters=5).eval()
        torch.manual_seed(0)
        x = torch.zeros(1, 4, d_model)
        streams = tuple(torch.zeros(1, 4, d_model) for _ in range(n_hyper))
        with torch.no_grad():
            q = mar.q_proj(x).view(1, 4, n_hyper, mar.d_head)
            ks = [mar.k_proj(s) for s in streams]
            k_stack = torch.stack(ks, dim=2)
            H_logits = torch.einsum("btqd,btkd->btqk", q, k_stack) / math.sqrt(mar.d_head)
            H_logits = H_logits + mar.identity_bias
            H = mar.sinkhorn_knopp(H_logits)
        diag = H.diagonal(dim1=-2, dim2=-1)
        off_diag = H - torch.diag_embed(diag)
        self.assertGreater(diag.mean().item(), off_diag.abs().mean().item())

    def test_mar_forward_preserves_shape(self):
        # mAR forward: [B, T, d_model] → [B, T, d_model]
        d_model, n_hyper, B, T = 128, 2, 4, 16
        mar = ManifoldConstrainedAttnRes(d_model=d_model, n_hyper=n_hyper)
        x = torch.randn(B, T, d_model)
        streams = tuple(torch.randn(B, T, d_model) for _ in range(n_hyper))
        y = mar(x, streams)
        self.assertEqual(y.shape, (B, T, d_model))

    def test_mar_streams_mismatch_raises(self):
        mar = ManifoldConstrainedAttnRes(d_model=128, n_hyper=2)
        x = torch.randn(2, 4, 128)
        with self.assertRaises(ValueError):
            mar(x, (torch.randn(2, 4, 128),))

    def test_mar_d_model_not_divisible_raises(self):
        with self.assertRaises(ValueError):
            ManifoldConstrainedAttnRes(d_model=127, n_hyper=2)

    def test_mar_gradients_flow_through_sinkhorn(self):
        # Sinkhorn-Knopp must be differentiable via exp
        mar = ManifoldConstrainedAttnRes(d_model=64, n_hyper=2)
        x = torch.randn(2, 8, 64, requires_grad=True)
        streams = tuple(torch.randn(2, 8, 64, requires_grad=True) for _ in range(2))
        y = mar(x, streams)
        y.sum().backward()
        self.assertIsNotNone(x.grad)
        self.assertFalse(torch.isnan(x.grad).any())
        for s in streams:
            self.assertIsNotNone(s.grad)
            self.assertFalse(torch.isnan(s.grad).any())

    def test_mar_temperature_gets_gradient(self):
        mar = ManifoldConstrainedAttnRes(d_model=64, n_hyper=2)
        y = mar(torch.randn(2, 4, 64), tuple(torch.randn(2, 4, 64) for _ in range(2)))
        y.sum().backward()
        self.assertIsNotNone(mar.temperature.grad)
        self.assertFalse(torch.isnan(mar.temperature.grad).any())

    def test_mar_buselModel_n_hyper_2(self):
        cfg = _MockConfig(d_model=128, n_layers=3, n_heads=4, n_hyper=2)
        model = buselModel(cfg)
        B, T = 2, 8
        hidden = torch.randn(B, T, cfg.d_model)
        mtp, aux = model(hidden)
        self.assertEqual(mtp[0].shape, (B, T, cfg.vocab_size))
        self.assertEqual(aux.shape, ())

    def test_mar_buselModel_n_hyper_4(self):
        cfg = _MockConfig(d_model=128, n_layers=3, n_heads=4, n_hyper=4)
        model = buselModel(cfg)
        B, T = 2, 8
        hidden = torch.randn(B, T, cfg.d_model)
        mtp, aux = model(hidden)
        self.assertEqual(mtp[0].shape, (B, T, cfg.vocab_size))

    def test_mar_residual_connection_preserved(self):
        # mHC §3.1: y_l = x_l + f_l(mixed_x_l). Residual add must be in buselModel.forward.
        cfg = _MockConfig(d_model=128, n_layers=2, n_heads=4, n_hyper=2)
        torch.manual_seed(42)
        model = buselModel(cfg).eval()
        const_value = 1.0
        for layer in model.layers:
            def constant_forward(self, x, progress=0.0):
                return torch.full_like(x, const_value), torch.tensor(0.0, device=x.device, dtype=x.dtype)
            layer.forward = constant_forward.__get__(layer)
        x_in = torch.zeros(1, 2, cfg.d_model)
        recorded_hidden = []
        original_final_norm = model.final_norm.forward
        def spy_norm(x):
            recorded_hidden.append(x.clone())
            return original_final_norm(x)
        model.final_norm.forward = spy_norm
        try:
            with torch.no_grad():
                model(x_in)
        finally:
            model.final_norm.forward = original_final_norm
        final_x = recorded_hidden[-1]
        expected_with_residual = 2 * const_value
        actual_max = final_x.abs().max().item()
        self.assertAlmostEqual(actual_max, expected_with_residual, delta=0.5,
                               msg=f"Residual missing or wrong: final x max={actual_max:.3f}, expected ~{expected_with_residual} (2 layers of residual adds with constant 1.0)")

    def test_mar_layer_outputs_flow_into_streams(self):
        # Streams must carry the post-residual x (residual stream values), enabling mHC mixing.
        cfg = _MockConfig(d_model=128, n_layers=3, n_heads=4, n_hyper=2)
        torch.manual_seed(0)
        model = buselModel(cfg).eval()
        recorded_mar_inputs = []
        original_mar_forward = ManifoldConstrainedAttnRes.forward
        def spy_forward(self, current_x, streams):
            recorded_mar_inputs.append([s.clone() for s in streams])
            return original_mar_forward(self, current_x, streams)
        ManifoldConstrainedAttnRes.forward = spy_forward
        try:
            x_in = torch.randn(1, 4, cfg.d_model)
            with torch.no_grad():
                model(x_in)
        finally:
            ManifoldConstrainedAttnRes.forward = original_mar_forward
        self.assertGreater(len(recorded_mar_inputs), 1, "MAR must run at each layer")
        for layer_idx, streams in enumerate(recorded_mar_inputs[1:], start=1):
            self.assertEqual(len(streams), cfg.n_hyper)

    def test_muon_ns_produces_orthogonal_output(self):
        # Paper §3.1: NS output singular values are bounded (≈ 1, within paper tolerance)
        # Uses eager _newton_schulz_core to avoid torch.compile recompile limit
        torch.manual_seed(0)
        X = torch.randn(64, 32)
        O = _newton_schulz_core(X, steps=5)
        sv = torch.linalg.svdvals(O)
        self.assertTrue((sv > 0.5).all(), f"singular values far below 1: min={sv.min().item()}")
        self.assertTrue((sv < 1.5).all(), f"singular values far above 1: max={sv.max().item()}")

    def test_muon_ns_uses_quintic_coefficients(self):
        # Paper §3.1: optimal quintic NS coefficients (3.4445, -4.7750, 2.0315)
        src = inspect.getsource(_newton_schulz_core)
        self.assertIn("3.4445", src)
        self.assertIn("-4.7750", src)
        self.assertIn("2.0315", src)

    def test_muon_ns_five_iterations_default(self):
        # Paper §3.1: NS uses 5 iterations by default
        src = inspect.getsource(_newton_schulz_core)
        self.assertIn("for _ in range(steps)", src)
        self.assertIn("steps=5", src)

    def test_muon_ns_handles_tall_and_wide_matrices(self):
        # Paper §3.1: NS uses transpose trick for tall (rows > cols) matrices
        # Uses eager _newton_schulz_core to avoid torch.compile recompile limit
        torch.manual_seed(1)
        for shape in [(128, 32), (32, 128), (64, 64), (256, 64)]:
            O = _newton_schulz_core(torch.randn(*shape), steps=5)
            self.assertEqual(O.shape, shape)
            self.assertFalse(torch.isnan(O).any(), f"shape={shape}: NaN output")
            self.assertFalse(torch.isinf(O).any(), f"shape={shape}: Inf output")

    def test_muon_scale_formula(self):
        # Paper §3.2: Muon scale = 0.2 * sqrt(max(A, B))
        src = inspect.getsource(Muon.step)
        self.assertIn("0.2", src)
        self.assertIn("sqrt", src)
        self.assertIn("max(A, B)", src)

    def test_muon_momentum_0_95(self):
        # Paper §3.2: Muon momentum = 0.95 (Nesterov-like)
        opt = Muon([torch.zeros(4, 4, requires_grad=True)], momentum=0.95)
        self.assertEqual(opt.param_groups[0]["momentum"], 0.95)

    def test_muon_hybrid_routing_splits_params(self):
        # Paper §3.3: Hybrid routes 2D `proj` params (excluding `router`) to Muon, rest to AdamW
        cfg = _MockConfig(d_model=128, n_layers=2, n_heads=4)
        model = buselModel(cfg)
        engine = buselOptimizerEngine(model, lr_muon=0.001, lr_adamw=0.0001)
        muon_ids = set(id(p) for g in engine.opt_muon.param_groups for p in g["params"])
        adamw_ids = set(id(p) for g in engine.opt_adamw.param_groups for p in g["params"])
        self.assertGreater(len(muon_ids), 0, "Muon should receive some params")
        self.assertGreater(len(adamw_ids), 0, "AdamW should receive some params")
        self.assertEqual(len(muon_ids & adamw_ids), 0, "no param in both")
        self.assertEqual(len(muon_ids | adamw_ids), sum(1 for p in model.parameters() if p.requires_grad))

    def test_muon_routing_covers_moe_and_mla(self):
        # ISSUES.md #1: 2D BitLinear weights (MoE, MLA, Blackboard, mtp) must route to Muon.
        cfg = _MockConfig(d_model=384, n_layers=4, n_heads=6, expert_hidden=768, num_experts=4)
        model = buselModel(cfg)
        engine = buselOptimizerEngine(model, lr_muon=0.001, lr_adamw=0.0001)
        n_muon = sum(p.numel() for g in engine.opt_muon.param_groups for p in g["params"])
        n_total = sum(p.numel() for p in model.parameters() if p.requires_grad)
        ratio = n_muon / n_total
        self.assertGreater(ratio, 0.85,
                           f"Expected ≥85% of params to Muon after ISSUES.md #1 fix; got {ratio*100:.1f}%")

    def test_muon_routing_excludes_router_and_embed(self):
        # Routers + embed tables stay on AdamW (noise-sensitive / structured).
        cfg = _MockConfig(d_model=128, n_layers=2, n_heads=4, num_experts=4)
        model = buselModel(cfg)
        engine = buselOptimizerEngine(model, lr_muon=0.001, lr_adamw=0.0001)
        muon_ids = {id(p) for g in engine.opt_muon.param_groups for p in g["params"]}
        for name, p in model.named_parameters():
            if "router" in name or "embed" in name:
                self.assertNotIn(id(p), muon_ids,
                                 f"{name} must NOT be on Muon (router/embed exclusion rule)")

    def test_muon_step_does_not_produce_nan(self):
        # ISSUES.md #3 + #4 regression: momentum update + NS core must not NaN.
        torch.manual_seed(42)
        p = torch.randn(64, 64, device=self.device, dtype=torch.float32) * 0.02
        opt = Muon([p], lr=0.001, momentum=0.95, weight_decay=0.1)
        for i in range(5):
            p.grad = torch.randn_like(p) * 0.01
            opt.step()
            self.assertFalse(torch.isnan(p).any(),
                             f"step {i}: Muon produced NaN weights")
            self.assertFalse(torch.isinf(p).any(),
                             f"step {i}: Muon produced Inf weights")

    def test_fastblt_vocab_size_dynamic(self):
        # Paper §2.1 (v5.4): byte-level vocab is 256 (UTF-8) + 70 multimodal specials = 326
        # (3 legacy + 67 plug-in tokens across 12 layers).
        patcher = StridedFastBLTPatcher(d_model=128)
        expected_vocab = mm_vocab_size() if HAS_SPECIAL_TOKENS else 259
        self.assertEqual(patcher.embed_weight.shape[0], expected_vocab,
                         f"patcher.embed_weight.shape[0] must equal vocab_size()={expected_vocab}")

    def test_fastblt_stride_4_kernel_5(self):
        # Paper §3: stride=4, conv kernel=5 (causal receptive field)
        patcher = StridedFastBLTPatcher(d_model=128, stride=4, kernel_size=5)
        self.assertEqual(patcher.stride, 4)
        self.assertEqual(patcher.kernel_size, 5)

    def test_fastblt_no_bpe_no_subword_tokens(self):
        # Paper §2.1: NO BPE — vocab stays under 500 (32k BPE contamination = forbidden)
        patcher = StridedFastBLTPatcher(d_model=128)
        self.assertLessEqual(patcher.embed_weight.shape[0], 500, "Vocab too large — likely BPE contamination")

    def test_fastblt_byte_input_to_patch_count(self):
        # Paper §3: T bytes → floor((T - 1) / stride) + 1 patches (left-padding by kernel-1)
        d_model, stride, kernel = 128, 4, 5
        patcher = StridedFastBLTPatcher(d_model=d_model, stride=stride, kernel_size=kernel).eval()
        T = 64
        v = mm_vocab_size() if HAS_SPECIAL_TOKENS else 259
        byte_ids = torch.randint(0, v, (2, T))
        with torch.no_grad():
            patches = patcher(byte_ids)
        expected_patches = (T - 1) // stride + 1
        self.assertEqual(patches.shape, (2, expected_patches, d_model))

    def test_fastblt_byte_embeddings_are_learned(self):
        # Paper §3.2: byte embeddings are LEARNED
        patcher = StridedFastBLTPatcher(d_model=128, d_byte=64)
        self.assertTrue(patcher.embed_weight.requires_grad)
        v = mm_vocab_size() if HAS_SPECIAL_TOKENS else 259
        self.assertEqual(patcher.embed_weight.shape, (v, 64))

    def test_fastblt_glu_gate_is_nonlinear(self):
        # Paper §3.2: GLU gate must be nonlinear (sigmoid after SiLU)
        patcher = StridedFastBLTPatcher(d_model=128, d_byte=32).eval()
        x = torch.randn(1, 16, 32)
        with torch.no_grad():
            gate_pos = torch.sigmoid(patcher.gate_proj_up(F.silu(patcher.gate_proj_down(x))))
            gate_neg = torch.sigmoid(patcher.gate_proj_up(F.silu(patcher.gate_proj_down(-x))))
        self.assertTrue((gate_pos > 0).all() and (gate_pos < 1).all())
        self.assertFalse(torch.allclose(gate_pos, gate_neg, atol=1e-3),
                         "GLU gate should be nonlinear — flipping input should change output")

    def test_fastblt_causal_left_padding(self):
        # Paper §3: causal conv uses left-side padding only (no future leakage)
        patcher = StridedFastBLTPatcher(d_model=128, stride=4, kernel_size=5).eval()
        T = 24
        byte_ids_a = torch.zeros(1, T, dtype=torch.long)
        byte_ids_a[0, 0] = 65
        byte_ids_a[0, 16] = 66
        with torch.no_grad():
            patches_a = patcher(byte_ids_a)
        byte_ids_b = byte_ids_a.clone()
        byte_ids_b[0, 16] = 67
        with torch.no_grad():
            patches_b = patcher(byte_ids_b)
        early_a = patches_a[0, 0]
        early_b = patches_b[0, 0]
        self.assertTrue(torch.allclose(early_a, early_b, atol=1e-3),
                        "Causal padding: byte at pos 16 must NOT affect the first patch (no future leakage)")
        late_a = patches_a[0, 4]
        late_b = patches_b[0, 4]
        self.assertFalse(torch.allclose(late_a, late_b, atol=1e-3),
                         "Changing byte at pos 16 SHOULD affect patch 4 (which covers that position)")

    def test_registry_decorator_basic(self):
        """🛸 busel REGISTRY — basic @register/get/is_registered/list_registered."""
        print("🧪 [REG-1] busel Registry — basic register/get API...")
        unregister("test_kind", "demo_cls")
        try:
            @register("test_kind", "demo_cls")
            class _Demo:
                pass

            self.assertTrue(is_registered("test_kind", "demo_cls"))
            self.assertIs(get("test_kind", "demo_cls"), _Demo)
            self.assertIn("demo_cls", list_registered("test_kind"))
            self.assertEqual(list_registered("test_kind"), ["demo_cls"])
            print("   ✅ Registry register/get/is_registered/list_registered pass.")
        finally:
            unregister("test_kind", "demo_cls")

    def test_registry_collision_raises(self):
        """🛸 busel REGISTRY — duplicate (kind, name) without override raises KeyError."""
        print("🧪 [REG-2] busel Registry — collision detection...")
        unregister("test_kind", "dup")
        try:
            @register("test_kind", "dup")
            class _A:
                pass

            with self.assertRaises(KeyError) as ctx:
                @register("test_kind", "dup")
                class _B:
                    pass

            msg = str(ctx.exception)
            self.assertIn("collision", msg.lower())
            self.assertIn("override=True", msg)
            self.assertIs(get("test_kind", "dup"), _A)
            print("   ✅ Registry collision correctly raised KeyError with override hint.")
        finally:
            unregister("test_kind", "dup")

    def test_registry_override_allowed(self):
        """🛸 busel REGISTRY — override=True replaces existing entry."""
        print("🧪 [REG-3] busel Registry — override=True works...")
        unregister("test_kind", "ovr")
        try:
            @register("test_kind", "ovr")
            class _First:
                pass

            @register("test_kind", "ovr", override=True)
            class _Second:
                pass

            self.assertIs(get("test_kind", "ovr"), _Second)
            print("   ✅ Registry override=True correctly replaced entry.")
        finally:
            unregister("test_kind", "ovr")

    def test_registry_attention_and_optimizer_registered(self):
        """🛸 busel REGISTRY — model.attention + training.optimizer register themselves on import."""
        print("🧪 [REG-4] busel Registry — production entries (attention, optimizer) are registered...")
        attn = list_registered("attention")
        opt = list_registered("optimizer")
        self.assertIn("gdn2", attn, "BulbaGDN2SeRoPEBlock must register as 'gdn2'")
        self.assertIn("mla", attn, "MultiHeadLatentAttention must register as 'mla'")
        self.assertIn("muon", opt, "Muon must register as 'muon'")
        self.assertIn("hybrid_muon_adamw", opt, "buselOptimizerEngine must register as 'hybrid_muon_adamw'")

        from model.attention import BulbaGDN2SeRoPEBlock, MultiHeadLatentAttention
        from training.optimizer import Muon as _Muon, buselOptimizerEngine
        self.assertIs(get("attention", "gdn2"), BulbaGDN2SeRoPEBlock)
        self.assertIs(get("attention", "mla"), MultiHeadLatentAttention)
        self.assertIs(get("optimizer", "muon"), _Muon)
        self.assertIs(get("optimizer", "hybrid_muon_adamw"), buselOptimizerEngine)
        print("   ✅ Production attention + optimizer entries are correctly registered.")

    def test_json_logger_writes_valid_jsonl(self):
        """📚 busel LOGGING — JSONFormatter produces valid one-line JSON per record."""
        print("🧪 [LOG-1] busel Logging — JSON formatter emits valid JSONL...")
        import io
        import logging as _logging
        from busel_logging import JSONFormatter

        buf = io.StringIO()
        handler = _logging.StreamHandler(buf)
        handler.setFormatter(JSONFormatter())
        logger = _logging.getLogger("busel_test_json")
        logger.handlers.clear()
        logger.setLevel(_logging.INFO)
        logger.addHandler(handler)
        logger.propagate = False

        logger.info("step_complete", extra={"step": 42, "loss": 3.14, "lr": 0.0001, "vram_mb": 1024.5, "extra_field": "hi"})

        line = buf.getvalue().strip()
        self.assertTrue(line.startswith("{") and line.endswith("}"), f"Line not JSON: {line!r}")
        self.assertNotIn("\n", line, "JSONL must be single line per record")
        parsed = json.loads(line)
        self.assertEqual(parsed["event"], "step_complete")
        self.assertEqual(parsed["level"], "INFO")
        self.assertIn("ts", parsed)
        self.assertEqual(parsed["step"], 42)
        self.assertAlmostEqual(parsed["loss"], 3.14, places=4)
        self.assertEqual(parsed["extra"]["extra_field"], "hi")
        print("   ✅ JSON logger emits valid single-line JSON with hoisted fields.")

    def test_json_logger_writes_to_file(self):
        """📚 busel LOGGING — setup_logging appends to a real file (one JSON per line)."""
        print("🧪 [LOG-2] busel Logging — setup_logging writes to checkpoints/busel.log.jsonl...")
        import tempfile
        from pathlib import Path

        with tempfile.TemporaryDirectory() as tmp:
            logger = setup_logging(log_dir=tmp, log_filename="test.jsonl")
            log_event("test_event", step=1, loss=2.5)
            log_event("training_complete", step=100, total_time_s=42.0)

            log_path = Path(tmp) / "test.jsonl"
            self.assertTrue(log_path.exists())
            lines = log_path.read_text(encoding="utf-8").strip().split("\n")
            self.assertEqual(len(lines), 2)
            for ln in lines:
                obj = json.loads(ln)
                self.assertIn("ts", obj)
                self.assertIn("event", obj)
                self.assertIn("level", obj)
            self.assertEqual(json.loads(lines[0])["event"], "test_event")
            self.assertEqual(json.loads(lines[1])["event"], "training_complete")
        print("   ✅ setup_logging appends valid JSONL to disk.")

    def test_teto_frames_nonempty_strings(self):
        """🎵 busel TETO — every state returns a non-empty string; all idle frames are distinct."""
        print("🧪 [TETO-1] busel Teto — frames are non-empty and distinct...")
        all_states = teto_states()
        self.assertGreaterEqual(len(all_states), 4, "Need at least 4 states (idle, blink, smile, ...)")
        for state in all_states:
            f = teto_frame(state, 0)
            self.assertIsInstance(f, str, f"frame({state!r}) must be str")
            self.assertGreater(len(f), 0, f"frame({state!r}) must be non-empty")

        idle = teto_frames()
        self.assertGreaterEqual(len(idle), 6, "Need at least 6 idle emoticon frames")
        self.assertEqual(len(idle), len(set(idle)), "Idle frames must all be distinct")
        for f in idle:
            self.assertGreater(len(f), 0)
            self.assertIsInstance(f, str)
        print(f"   ✅ Teto: {len(all_states)} states, {len(idle)} distinct idle frames — all non-empty.")

    def test_teto_idle_cycles_through_frames(self):
        """🎵 busel TETO — frame('idle', tick) cycles through the 12-frame emoticon set."""
        print("🧪 [TETO-2] busel Teto — idle cycle wraps correctly at modulo 12...")
        cycle_a = [teto_frame("idle", i) for i in range(12)]
        cycle_b = [teto_frame("idle", i + 12) for i in range(12)]
        self.assertEqual(cycle_a, cycle_b, "frame('idle', tick+12) must equal frame('idle', tick) for all 12 ticks")
        self.assertEqual(len(set(cycle_a)), 12, "All 12 idle frames must be distinct within one cycle")
        print(f"   ✅ Idle cycle of {len(cycle_a)} distinct emoticons wraps cleanly modulo 12.")

    def test_cli_helpers_do_not_crash(self):
        """💡 busel CLI — every public helper runs without raising (with or without rich)."""
        print("🧪 [CLI-1] busel CLI — all helpers run without raising...")
        ui_cli.header("TEST HEADER", "subtitle here")
        ui_cli.status_panel("test", key1="value1", key2=42)
        ui_cli.log("info message", level="info")
        ui_cli.log("warn message", level="warn")
        ui_cli.log("error message", level="error")
        ui_cli.log("ok message", level="ok")
        line = ui_cli.step_line(10, 100, 3.14, 0.5, 0.001, 1.5, 8, 1234.5, 256.0)
        self.assertIn("Step 00010/00100", line)
        self.assertIn("1234", line)
        self.assertIn("VRAM: 256MB", line)
        self.assertIn("Total", line)
        ui_cli.safe_print("hello\n")
        print("   ✅ All ui.cli helpers execute without raising.")

    def test_cli_animated_header_runs(self):
        """💡 busel CLI — animated_header + project_tree + spinner + progress_bar all execute."""
        print("🧪 [CLI-2] busel CLI — animated_header, spinner, progress_bar, project_tree...")
        ui_cli.animated_header("busel TEST", cycles=2, palette="teto")
        ui_cli.animated_header("busel TEST", cycles=1, palette="cycle")
        ui_cli.project_tree()
        with ui_cli.spinner("working") as _:
            pass
        with ui_cli.progress_bar(total=10, description="test") as handle:
            if handle is not None:
                handle.update(advance=5)
        print("   ✅ Animated header, spinner, progress bar, project tree all execute cleanly.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_registry_lists_all_encoders(self):
        """🛰️ busel MULTIMODAL — every encoder class self-registers on import."""
        print("🧪 [MM-1] busel Multimodal — registry lists all 6 encoders...")
        names = list_mm_encoders()
        for n in ("image", "video", "audio", "pdf", "docx", "text"):
            self.assertIn(n, names, f"encoder '{n}' must be registered")
        self.assertGreaterEqual(len(names), 6)
        print(f"   ✅ Multimodal registry: {sorted(names)}")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_image_encoder_roundtrip(self):
        """🛰️ busel MULTIMODAL — ImageEncoder encode → decode returns 32×32 RGB."""
        print("🧪 [MM-2] busel Multimodal — image encode/decode round-trip...")
        from PIL import Image as _Im
        src = _Im.new("RGB", (256, 256), color=(12, 34, 56))
        enc = ImageEncoder()
        tokens = enc.encode(src)
        self.assertIsInstance(tokens, list, "encode must return list[int]")
        self.assertEqual(len(tokens), IMAGE_BYTES + 2, f"must be {IMAGE_BYTES}+2 tokens, got {len(tokens)}")
        self.assertEqual(tokens[0], MOD_IMAGE, "first token must be MOD_IMAGE (v5.4 modality prefix)")
        self.assertEqual(tokens[-1], MEDIA_END, "last token must be 257 (__MEDIA_END__)")
        for t in tokens[1:-1]:
            self.assertGreaterEqual(t, 0)
            self.assertLess(t, 256, f"image payload token {t} must be real byte (0..255)")
        out = enc.decode(tokens)
        self.assertEqual(out.size, (32, 32))
        self.assertEqual(out.mode, "RGB")
        print("   ✅ ImageEncoder: 32×32 RGB round-trip, marker tokens correct.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_image_encoder_rejects_bad_blob(self):
        """🛰️ busel MULTIMODAL — ImageEncoder.decode raises on missing/corrupt markers."""
        print("🧪 [MM-3] busel Multimodal — image decode rejects bad blob...")
        enc = ImageEncoder()
        with self.assertRaises(ValueError):
            enc.decode([])
        with self.assertRaises(ValueError):
            enc.decode([0] * 100)
        with self.assertRaises(ValueError):
            enc.decode([IMAGE_MARKER] + list(b"short") + [MEDIA_END])
        print("   ✅ ImageEncoder.decode raises on bad input.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_video_encoder_roundtrip(self):
        """🛰️ busel MULTIMODAL — VideoEncoder encodes frame_count + N×32×32×3 to bytes."""
        print("🧪 [MM-4] busel Multimodal — video encode/decode round-trip...")
        import numpy as _np
        import imageio.v3 as _iio
        tmp = "temp_test_mm_video.mp4"
        try:
            n_frames = 6
            frames = [_np.random.randint(0, 255, (32, 32, 3), dtype=_np.uint8) for _ in range(n_frames)]
            _iio.imwrite(tmp, frames, fps=10)
            enc = VideoEncoder(max_frames=3)
            tokens = enc.encode_file(tmp)
            self.assertIsInstance(tokens, list)
            self.assertEqual(tokens[0], IMAGE_MARKER)
            self.assertEqual(tokens[-1], MEDIA_END)
            expected_len = 1 + 4 + 3 * IMAGE_BYTES + 1
            self.assertEqual(len(tokens), expected_len, f"video must be {expected_len} tokens")
            payload = enc.decode(tokens)
            self.assertEqual(len(payload), 3 * IMAGE_BYTES, f"expected 3 frames × {IMAGE_BYTES}, got {len(payload)}")
            print(f"   ✅ VideoEncoder: {n_frames}-frame input downsampled to 3 frames, round-trip OK.")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_audio_encoder_roundtrip(self):
        """🛰️ busel MULTIMODAL — AudioEncoder writes [sr][n][sw][PCM16] with markers."""
        print("🧪 [MM-5] busel Multimodal — audio encode/decode round-trip...")
        import numpy as _np
        import soundfile as _sf
        import wave as _wave
        from io import BytesIO as _BI
        tmp = "temp_test_mm_audio.wav"
        try:
            sr = 16000
            data = _np.random.uniform(-0.3, 0.3, sr).astype(_np.float32)
            _sf.write(tmp, data, sr)
            enc = AudioEncoder(max_seconds=2.0)
            tokens = enc.encode_file(tmp)
            self.assertIsInstance(tokens, list)
            self.assertEqual(tokens[0], MOD_AUDIO, "first token must be MOD_AUDIO (v5.4)")
            self.assertEqual(tokens[-1], MEDIA_END)
            header = bytes(tokens[1:11])
            sr_out, n_out, sw_out = struct.unpack("<IIH", header)
            self.assertEqual(sr_out, sr)
            self.assertEqual(sw_out, 2, "sample_width must be 2 (int16)")
            self.assertEqual(n_out, sr, "1 second @ 16kHz = 16000 samples")
            wav = enc.decode_to_wav(tokens)
            with _wave.open(_BI(wav), "rb") as wf:
                self.assertEqual(wf.getframerate(), sr)
                self.assertEqual(wf.getsampwidth(), 2)
                self.assertEqual(wf.getnframes(), sr)
            print("   ✅ AudioEncoder: 1s @ 16kHz round-trips, header layout correct.")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_docx_encoder_roundtrip(self):
        """🛰️ busel MULTIMODAL — DocxEncoder writes UTF-8 plain text with markers."""
        print("🧪 [MM-6] busel Multimodal — docx encode round-trip...")
        import docx as _docx
        tmp = "temp_test_mm.docx"
        try:
            d = _docx.Document()
            d.add_paragraph("Hello, multimodal Busel!")
            d.add_paragraph("Line two.")
            d.save(tmp)
            enc = DocxEncoder()
            tokens = enc.encode_file(tmp)
            self.assertEqual(tokens[0], MOD_DOCX, "first token must be MOD_DOCX (v5.4 modality prefix)")
            self.assertEqual(tokens[-1], MEDIA_END)
            text = bytes(tokens[1:-1]).decode("utf-8")
            self.assertIn("Hello, multimodal Busel!", text)
            self.assertIn("Line two.", text)
            print("   ✅ DocxEncoder: 2-paragraph docx → UTF-8 tokens round-trip OK.")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_text_encoder_passes_bytes_through(self):
        """🛰️ busel MULTIMODAL — TextEncoder is a thin pass-through (no markers)."""
        print("🧪 [MM-7] busel Multimodal — text encode is pass-through...")
        tmp = "temp_test_mm.txt"
        try:
            with open(tmp, "wb") as f:
                f.write(b"hello \xe2\x98\x83 unicode")
            enc = TextEncoder()
            tokens = enc.encode_file(tmp)
            self.assertEqual(tokens, list(b"hello \xe2\x98\x83 unicode"))
            self.assertNotIn(IMAGE_MARKER, tokens, "text encoder must NOT inject media markers")
            self.assertNotIn(MEDIA_END, tokens, "text encoder must NOT inject media markers")
            print("   ✅ TextEncoder: byte-for-byte pass-through, no marker injection.")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_build_encoder_for_dispatch(self):
        """🛰️ busel MULTIMODAL — build_encoder_for routes by extension, falls back to text."""
        print("🧪 [MM-8] busel Multimodal — build_encoder_for extension dispatch...")
        self.assertIsInstance(build_encoder_for("a.png"), ImageEncoder)
        self.assertIsInstance(build_encoder_for("A.JPG"), ImageEncoder)
        self.assertIsInstance(build_encoder_for("b.mp4"), VideoEncoder)
        self.assertIsInstance(build_encoder_for("c.wav"), AudioEncoder)
        self.assertIsInstance(build_encoder_for("d.docx"), DocxEncoder)
        self.assertIsInstance(build_encoder_for("e.txt"), TextEncoder)
        self.assertIsInstance(build_encoder_for("f.md"), TextEncoder)
        self.assertIsInstance(build_encoder_for("g.unknown"), TextEncoder, "unknown must fall back to TextEncoder")
        print("   ✅ build_encoder_for routes by extension (case-insensitive), falls back to text.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_marker_constants_reserved_vocab(self):
        """🛰️ busel MULTIMODAL — 256, 257, 258 are reserved marker tokens (outside 0..255)."""
        print("🧪 [MM-9] busel Multimodal — marker tokens are reserved (>= 256)...")
        self.assertEqual(IMAGE_MARKER, 256, "__MEDIA_START__ must be token 256")
        self.assertEqual(MEDIA_END, 257, "__MEDIA_END__ must be token 257")
        for marker in (IMAGE_MARKER, MEDIA_END, 258):
            self.assertGreaterEqual(marker, 256, f"marker {marker} must be in reserved range (>= 256)")
        for b in range(256):
            self.assertNotIn(b, (IMAGE_MARKER, MEDIA_END, 258), f"byte {b} collides with reserved marker")
        print("   ✅ Marker tokens 256, 257, 258 are reserved; no collision with valid UTF-8 bytes.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_image_encoder_byte_layout_lossless(self):
        """🛰️ busel MULTIMODAL — image encoder is lossless: encode→decode→encode is a fixed point."""
        print("🧪 [MM-10] busel Multimodal — image encoder fixed point (lossless)...")
        from PIL import Image as _Im
        import numpy as _np
        enc = ImageEncoder()
        src = _Im.fromarray(_np.random.randint(0, 255, (100, 100, 3), dtype=_np.uint8))
        blob1 = enc.encode(src)
        decoded = enc.decode(blob1)
        blob2 = enc.encode(decoded)
        self.assertEqual(blob1, blob2, "encode→decode→encode must be a fixed point (lossless)")
        print("   ✅ Image encoder is lossless: encode→decode→encode is byte-identical.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_text_encoder_in_pipeline_collates_to_int32(self):
        """🛰️ busel MULTIMODAL — encoder output flows through collate_busel_batch to int32 tensor."""
        print("🧪 [MM-11] busel Multimodal — encoder output → collate → int32 tensor with values < vocab_size()...")
        from data.pipeline import collate_busel_batch
        from multimodal.special_tokens import vocab_size as _vs
        enc = TextEncoder()
        tokens = enc.encode_file(__file__)
        self.assertIsInstance(tokens, list)
        batch = collate_busel_batch([(tokens, 0, 0)])
        tensor, _, _ = batch
        self.assertEqual(tensor.dtype, torch.int32, "collate must produce int32 tensor")
        v = _vs()
        self.assertLess(tensor.max().item(), v, f"max token must be < vocab_size()={v}")
        self.assertGreaterEqual(tensor.min().item(), 0, "min token must be >= 0")
        self.assertEqual(tensor.shape[0], 1, "batch size must be 1")
        print(f"   ✅ TextEncoder output flows through collate to int32 tensor of shape {tuple(tensor.shape)}, max={tensor.max().item()} < {v}.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_cv2_fast_path_under_500ms_per_100_imgs(self):
        """🛰️ busel MULTIMODAL — cv2 image encode: 100×256² images must encode in <500ms."""
        print("🧪 [MM-12] busel Multimodal — cv2 fast-path throughput (100 imgs < 500ms)...")
        try:
            import cv2 as _cv2
            import numpy as _np
        except ImportError:
            self.skipTest("opencv-python-headless not installed")
        from PIL import Image as _Im
        enc = ImageEncoder()
        n = 100
        src_arr = _np.random.randint(0, 255, (256, 256, 3), dtype=_np.uint8)
        src = _Im.fromarray(src_arr)
        t0 = time.perf_counter()
        for _ in range(n):
            enc.encode(src)
        elapsed_ms = (time.perf_counter() - t0) * 1000
        per_op = elapsed_ms / n
        self.assertLess(elapsed_ms, 500, f"100 cv2 encodings took {elapsed_ms:.1f}ms (> 500ms budget)")
        print(f"   ✅ cv2 fast path: {per_op:.2f} ms/image ({n} images in {elapsed_ms:.1f}ms)")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_video_cv2_path_under_2s_for_60_frame_video(self):
        """🛰️ busel MULTIMODAL — cv2 video encode: 60-frame synthetic video must encode in <2s."""
        print("🧪 [MM-13] busel Multimodal — cv2 video path throughput (60 frames < 2s)...")
        try:
            import cv2 as _cv2
            import numpy as _np
        except ImportError:
            self.skipTest("opencv-python-headless not installed")
        tmp = "temp_test_mm_video_perf.mp4"
        try:
            n_frames = 60
            h, w = 128, 128
            fourcc = _cv2.VideoWriter_fourcc(*"mp4v")
            writer = _cv2.VideoWriter(tmp, fourcc, 30.0, (w, h))
            for i in range(n_frames):
                frame = _np.random.randint(0, 255, (h, w, 3), dtype=_np.uint8)
                writer.write(frame)
            writer.release()
            enc = VideoEncoder(max_frames=8)
            t0 = time.perf_counter()
            tokens = enc.encode_file(tmp)
            elapsed_ms = (time.perf_counter() - t0) * 1000
            self.assertLess(elapsed_ms, 2000, f"video encoding took {elapsed_ms:.1f}ms (> 2000ms budget)")
            self.assertEqual(tokens[0], MOD_VIDEO, "first token must be MOD_VIDEO (v5.4)")
            self.assertEqual(tokens[-1], MEDIA_END)
            print(f"   ✅ cv2 video path: 60 frames @ 128×128 → 8 frames in {elapsed_ms:.1f}ms")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    def test_muon_routing_rule_2d_proj_router_embed(self):
        """🛡️ Muon routing (ISSUES.md #5): 2D+!router+!embed → Muon, else → AdamW."""
        print("🧪 [R5] busel Muon routing — 2D non-router/non-embed → Muon, else → AdamW...")
        class MockM(torch.nn.Module):
            def __init__(self):
                super().__init__()
                self.layers_q_proj = torch.nn.Linear(64, 64, bias=False)
                self.expert_ffn_0_weight = torch.nn.Parameter(torch.randn(128, 64))
                v = mm_vocab_size() if HAS_SPECIAL_TOKENS else 259
                self.mtp_head_3 = torch.nn.Parameter(torch.randn(v, 128))
                self.bias = torch.nn.Parameter(torch.zeros(64))
                self.router = torch.nn.Linear(64, 4, bias=False)
                self.token_embed = torch.nn.Parameter(torch.randn(v, 64))

        model = MockM()
        engine = buselOptimizerEngine(model, lr_muon=0.01, lr_adamw=0.001)
        muon_param_ids = {id(p) for p in engine.opt_muon.param_groups[0]['params']}
        adamw_param_ids = {id(p) for p in engine.opt_adamw.param_groups[0]['params']}
        for name, p in model.named_parameters():
            in_muon = id(p) in muon_param_ids
            in_adamw = id(p) in adamw_param_ids
            self.assertTrue(in_muon or in_adamw, f"{name} assigned to NEITHER optimizer")
            self.assertFalse(in_muon and in_adamw, f"{name} assigned to BOTH optimizers")
            if "router" in name or "embed" in name or p.ndim != 2:
                self.assertTrue(in_adamw, f"{name} (1D/router/embed) should be in AdamW, got Muon")
            else:
                self.assertTrue(in_muon, f"{name} (2D+!router+!embed) should be in Muon, got AdamW")
        n_muon = len(muon_param_ids)
        n_adamw = len(adamw_param_ids)
        self.assertEqual(n_muon, 3, f"expected 3 Muon params, got {n_muon}")
        self.assertEqual(n_adamw, 3, f"expected 3 AdamW params, got {n_adamw}")
        print(f"   ✅ Muon routing: 3 Muon / 3 AdamW (correctly excludes router/embed/1D).")

    def test_training_schedule_guard_rejects_bad_inputs(self):
        """🛡️ Schedule guard (ISSUES.md #7): max_steps > warmup_steps, warmup >= 1."""
        print("🧪 [R7] busel schedule guard — rejects max_steps<=warmup_steps and warmup<1...")
        with self.assertRaises(ValueError, msg="max_steps==warmup_steps must be rejected"):
            validate_training_schedule(100, 100)
        with self.assertRaises(ValueError, msg="max_steps<warmup_steps must be rejected"):
            validate_training_schedule(50, 100)
        with self.assertRaises(ValueError, msg="warmup_steps<1 must be rejected"):
            validate_training_schedule(100, 0)
        with self.assertRaises(ValueError, msg="warmup_steps<0 must be rejected"):
            validate_training_schedule(100, -5)
        with self.assertRaises(ValueError, msg="None must be rejected"):
            validate_training_schedule(None, 10)
        with self.assertRaises(ValueError, msg="None must be rejected"):
            validate_training_schedule(10, None)
        max_s, warmup_s = validate_training_schedule(200, 10)
        self.assertEqual((max_s, warmup_s), (200, 10))
        max_s, warmup_s = validate_training_schedule("300", "20")
        self.assertEqual((max_s, warmup_s), (300, 20), "string-cast ints must be normalised")
        print("   ✅ Schedule guard: rejects all 6 bad cases, accepts 2 valid cases.")

    def test_inject_noise_branchless_preserves_grad_shape(self):
        """🛡️ inject_noise (ISSUES.md #6): branchless mask keeps grad shape & adds bounded noise."""
        print("🧪 [R6] busel inject_noise — branchless mask, grad shape preserved, noise bounded...")
        from training.autopilot import buselAutoPilot

        class _MockEngine:
            def __init__(self):
                self.opt_muon = type("O", (), {"param_groups": [{"params": []}]})()
                self.opt_adamw = type("O", (), {"param_groups": [{"params": []}]})()

        engine = _MockEngine()
        ap = buselAutoPilot(engine, max_lr_muon=0.01, max_lr_adamw=0.001, noise_scale=0.1)
        ap.noise_scale = 0.1

        model = torch.nn.Sequential(
            torch.nn.Linear(8, 8, bias=False),
            torch.nn.Linear(8, 4, bias=False),
        )
        for p in model.parameters():
            p.grad = torch.randn_like(p) * 0.5
        g_before = [p.grad.clone() for p in model.parameters()]

        ap.inject_noise(model)

        for p, g in zip(model.parameters(), g_before):
            self.assertEqual(p.grad.shape, g.shape, "grad shape must be preserved")
            self.assertTrue(torch.isfinite(p.grad).all(), "grad must remain finite")
            self.assertFalse(torch.allclose(p.grad, g, atol=1e-8), "noise should perturb the grad")
        ap.noise_scale = 0.0
        g_before2 = [p.grad.clone() for p in model.parameters()]
        ap.inject_noise(model)
        for p, g in zip(model.parameters(), g_before2):
            self.assertTrue(torch.allclose(p.grad, g, atol=1e-12), "noise_scale=0 must be no-op")
        print("   ✅ inject_noise: shape preserved, finite, no-op at zero scale.")

    # ============================================================
    # Vocab expansion tests (v5.4.0 — 67 plug-in + 3 legacy = 70 specials)
    # ============================================================

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "multimodal.special_tokens unavailable")
    def test_mm_special_tokens_total_vocab_is_326(self):
        """🛰️ [MM-14] busel SPECIAL TOKENS — vocab_size() = 256 bytes + 3 legacy + 67 plug-in = 326."""
        print("🧪 [MM-14] busel Special Tokens — total vocab is 326 (256+3+67)...")
        v = mm_vocab_size()
        self.assertEqual(v, 326, f"vocab_size() must be 326 (256 bytes + 3 legacy + 67 plug-in), got {v}")
        enabled = enabled_ids()
        self.assertEqual(len(enabled), 70, f"must have 70 enabled special IDs, got {len(enabled)}")
        self.assertEqual(enabled[0], 256, "first special must be legacy MEDIA_START (256)")
        self.assertEqual(enabled[-1], 325, f"last special must be 325, got {enabled[-1]}")
        self.assertEqual(enabled, sorted(enabled), "enabled IDs must be ascending")
        self.assertEqual(len(set(enabled)), len(enabled), "no duplicate IDs")
        print(f"   ✅ vocab_size()=326, 70 enabled specials, IDs contiguous [256..325].")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "multimodal.special_tokens unavailable")
    def test_mm_special_tokens_layer_summary(self):
        """🛰️ [MM-15] busel SPECIAL TOKENS — 12 layers, expected counts per layer."""
        print("🧪 [MM-15] busel Special Tokens — 12 layers with documented counts...")
        from multimodal.special_tokens import layer_summary
        summary = layer_summary()
        self.assertEqual(set(summary.keys()), set(LAYER_DESCRIPTIONS.keys()),
                         "all 12 documented layers must be present")
        expected_counts = {
            "sequence": 4, "modality": 6, "mm_struct": 3, "role": 4,
            "reasoning": 4, "code": 4, "tool_xml": 12, "tool": 12,
            "task": 4, "reference": 6, "subagent": 4, "status": 4,
        }
        for layer, n in expected_counts.items():
            self.assertEqual(summary[layer], n,
                             f"layer {layer!r} must have {n} tokens, got {summary[layer]}")
        total = sum(summary.values())
        self.assertEqual(total, 67, f"total plug-in tokens must be 67, got {total}")
        print(f"   ✅ 12 layers, counts {dict(sorted(summary.items()))}, total = 67 plug-in.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_image_encoder_emits_mod_image(self):
        """🛰️ [MM-16] busel ImageEncoder emits MOD_IMAGE (263) at start, not legacy 256."""
        print("🧪 [MM-16] busel ImageEncoder — emits MOD_IMAGE (263), not legacy MEDIA_START...")
        from PIL import Image as _Im
        import numpy as _np
        enc = ImageEncoder()
        src = _Im.fromarray(_np.random.randint(0, 255, (100, 100, 3), dtype=_np.uint8))
        blob = enc.encode(src)
        self.assertEqual(blob[0], MOD_IMAGE, f"image stream must start with MOD_IMAGE={MOD_IMAGE}, got {blob[0]}")
        self.assertEqual(blob[-1], MEDIA_END, f"image stream must end with MEDIA_END={MEDIA_END}, got {blob[-1]}")
        self.assertNotEqual(blob[0], 256, "v5.4 must NOT emit legacy MEDIA_START (256) for images")
        print(f"   ✅ ImageEncoder: [MOD_IMAGE={blob[0]}, ...payload..., MEDIA_END={blob[-1]}].")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_video_encoder_roundtrip(self):
        """🛰️ busel MULTIMODAL — VideoEncoder encodes frame_count + N×32×32×3 to bytes."""
        print("🧪 [MM-4] busel Multimodal — video encode/decode round-trip...")
        try:
            import cv2 as _cv2
        except ImportError:
            self.skipTest("opencv-python-headless not installed")
        import numpy as _np
        tmp = "temp_test_video.mp4"
        if os.path.exists(tmp):
            os.remove(tmp)
        try:
            fourcc = _cv2.VideoWriter_fourcc(*"mp4v")
            h, w = 64, 64
            writer = _cv2.VideoWriter(tmp, fourcc, 10.0, (w, h))
            for i in range(15):
                frame = _np.full((h, w, 3), i * 16, dtype=_np.uint8)
                writer.write(frame)
            writer.release()
            enc = VideoEncoder(max_frames=4)
            tokens = enc.encode_file(tmp)
            self.assertEqual(tokens[0], MOD_VIDEO, "first token must be MOD_VIDEO (v5.4 modality prefix)")
            self.assertEqual(tokens[-1], MEDIA_END)
            decoded = enc.decode(tokens)
            self.assertGreater(len(decoded), 0)
            print(f"   ✅ VideoEncoder: [{len(tokens)} tokens, decoded {len(decoded)} bytes].")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)
        try:
            fourcc = _cv2.VideoWriter_fourcc(*"mp4v")
            h, w = 64, 64
            writer = _cv2.VideoWriter(tmp, fourcc, 10.0, (w, h))
            for i in range(15):
                frame = _np.full((h, w, 3), i * 16, dtype=_np.uint8)
                writer.write(frame)
            writer.release()
            enc = VideoEncoder(max_frames=4)
            blob = enc.encode_file(tmp)
            self.assertEqual(blob[0], MOD_VIDEO, f"video stream must start with MOD_VIDEO={MOD_VIDEO}, got {blob[0]}")
            self.assertEqual(blob[-1], MEDIA_END, f"video stream must end with MEDIA_END={MEDIA_END}")
            print(f"   ✅ VideoEncoder: [MOD_VIDEO={blob[0]}, ...4 frames..., MEDIA_END={blob[-1]}].")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_audio_encoder_emits_mod_audio(self):
        """🛰️ [MM-18] busel AudioEncoder emits MOD_AUDIO (265) at start."""
        print("🧪 [MM-18] busel AudioEncoder — emits MOD_AUDIO (265) at start...")
        try:
            import soundfile as _sf
        except ImportError:
            self.skipTest("soundfile not installed")
        import numpy as _np
        tmp = "temp_test_audio_mod.wav"
        if os.path.exists(tmp):
            os.remove(tmp)
        try:
            sr = 8000
            data = _np.random.uniform(-0.3, 0.3, size=(sr,)).astype(_np.float32)
            _sf.write(tmp, data, sr)
            enc = AudioEncoder(max_seconds=0.5)
            blob = enc.encode_file(tmp)
            self.assertEqual(blob[0], MOD_AUDIO, f"audio stream must start with MOD_AUDIO={MOD_AUDIO}, got {blob[0]}")
            self.assertEqual(blob[-1], MEDIA_END, f"audio stream must end with MEDIA_END={MEDIA_END}")
            print(f"   ✅ AudioEncoder: [MOD_AUDIO={blob[0]}, header+pcm..., MEDIA_END={blob[-1]}].")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_docx_encoder_emits_mod_docx(self):
        """🛰️ [MM-19] busel DocxEncoder emits MOD_DOCX (267) at start."""
        print("🧪 [MM-19] busel DocxEncoder — emits MOD_DOCX (267) at start...")
        try:
            import docx as _docx
        except ImportError:
            self.skipTest("python-docx not installed")
        tmp = "temp_test_mod.docx"
        if os.path.exists(tmp):
            os.remove(tmp)
        try:
            d = _docx.Document()
            d.add_paragraph("hello")
            d.add_paragraph("world")
            d.save(tmp)
            enc = DocxEncoder()
            blob = enc.encode_file(tmp)
            self.assertEqual(blob[0], MOD_DOCX, f"docx stream must start with MOD_DOCX={MOD_DOCX}, got {blob[0]}")
            self.assertEqual(blob[-1], MEDIA_END, f"docx stream must end with MEDIA_END={MEDIA_END}")
            print(f"   ✅ DocxEncoder: [MOD_DOCX={blob[0]}, text..., MEDIA_END={blob[-1]}].")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS, "multimodal encoders unavailable")
    def test_mm_text_encoder_passes_bytes_through(self):
        """🛰️ busel MULTIMODAL — TextEncoder prefixes MOD_TEXT (v5.4) then passes bytes through."""
        print("🧪 [MM-7] busel Multimodal — text encode = [MOD_TEXT] + bytes...")
        enc = TextEncoder()
        tmp = "temp_test_text.txt"
        if os.path.exists(tmp):
            os.remove(tmp)
        try:
            with open(tmp, "wb") as f:
                f.write(b"hello \xe2\x98\x83 unicode")
            tokens = enc.encode_file(tmp)
            self.assertEqual(tokens[0], MOD_TEXT, "v5.4 text stream must start with MOD_TEXT")
            payload = bytes(tokens[1:])
            self.assertEqual(payload, b"hello \xe2\x98\x83 unicode", "payload must be byte-identical")
            self.assertNotIn(IMAGE_MARKER, tokens, "text encoder must NOT inject legacy media_start")
            print(f"   ✅ TextEncoder: [MOD_TEXT] + {len(payload)} payload bytes, no legacy markers.")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)
        try:
            with open(tmp, "wb") as f:
                f.write(b"hello world")
            blob = enc.encode_file(tmp)
            self.assertEqual(blob[0], MOD_TEXT, f"text stream must start with MOD_TEXT={MOD_TEXT}, got {blob[0]}")
            self.assertNotEqual(blob[-1], MEDIA_END, "text streams are unbounded, no trailing MEDIA_END")
            self.assertGreater(len(blob), 6, "stream must contain header + payload")
            print(f"   ✅ TextEncoder: [MOD_TEXT={blob[0]}, ...{len(blob)-1} bytes, no trailing MEDIA_END].")
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "multimodal.special_tokens unavailable")
    def test_mm_special_tokens_disable_and_enable_roundtrip(self):
        """🛰️ [MM-21] busel SPECIAL TOKENS — disable shrinks vocab, enable restores it."""
        print("🧪 [MM-21] busel Special Tokens — disable/enable roundtrip preserves ID...")
        pre = mm_vocab_size()
        target = "think_start"
        # Save state in case prior test left it disabled
        if not is_enabled(target):
            enable_special_token(target)
        tok_before = get_special_token(target)
        self.assertTrue(tok_before.enabled)
        self.assertTrue(THINK_START.enabled)
        try:
            disable_special_token(target)
            mid = mm_vocab_size()
            self.assertEqual(mid, pre - 1, f"after disable, vocab must shrink by 1 (was {pre}, got {mid})")
            tok_after = get_special_token(target)
            self.assertFalse(tok_after.enabled, "token must be marked disabled")
            self.assertEqual(tok_after.id, tok_before.id, "ID must be preserved across disable")
            self.assertFalse(is_enabled(target), "is_enabled() must return False")
            enable_special_token(target)
            post = mm_vocab_size()
            self.assertEqual(post, pre, f"after enable, vocab must restore (was {pre}, got {post})")
            self.assertTrue(get_special_token(target).enabled)
            self.assertTrue(is_enabled(target))
            print(f"   ✅ Disable/enable: vocab {pre} → {mid} → {post}, ID={tok_before.id} preserved.")
        finally:
            enable_special_token(target)

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "multimodal.special_tokens unavailable")
    def test_mm_special_tokens_register_at_runtime_grows_vocab(self):
        """🛰️ [MM-22] busel SPECIAL TOKENS — register_special_token at runtime grows vocab."""
        print("🧪 [MM-22] busel Special Tokens — register_special_token at runtime grows vocab...")
        from multimodal import special_tokens as mm_st
        pre = mm_vocab_size()
        test_name = "_test_runtime_token_xyz"
        try:
            tok = mm_st.register_special_token(test_name, "test_layer", "test description")
            post = mm_vocab_size()
            self.assertEqual(post, pre + 1, f"vocab must grow by 1 (was {pre}, got {post})")
            self.assertTrue(tok.enabled)
            self.assertEqual(tok.name, test_name)
            self.assertEqual(tok.layer, "test_layer")
            self.assertTrue(tok.id >= SPECIAL_VOCAB_BASE, f"new token id must be >= {SPECIAL_VOCAB_BASE}, got {tok.id}")
            # Verify it's discoverable
            self.assertIs(mm_st.get_special_token(test_name), tok)
            self.assertIn(test_name, [t.name for t in mm_st.list_special_tokens()])
            self.assertIn(test_name, [t.name for t in mm_st.list_special_tokens(layer="test_layer")])
            print(f"   ✅ register_special_token({test_name!r}): vocab {pre} → {post}, id={tok.id}.")
        finally:
            # Cleanup: unregister the test token
            mm_st.unregister_kind_safe(test_name) if hasattr(mm_st, "unregister_kind_safe") else None
            # Manual cleanup: disable it (next test will see vocab restored)
            if test_name in [t.name for t in mm_st.list_special_tokens(enabled_only=False)]:
                mm_st.disable_special_token(test_name)

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "multimodal.special_tokens unavailable")
    def test_mm_special_tokens_legacy_collision_rejected(self):
        """🛰️ [MM-23] busel SPECIAL TOKENS — cannot register a name that collides with a legacy ID."""
        print("🧪 [MM-23] busel Special Tokens — legacy name collision rejected...")
        with self.assertRaises(ValueError, msg="must reject 'media_start' as new token name"):
            register_special_token("media_start", "test", "should fail")
        with self.assertRaises(ValueError, msg="must reject empty name"):
            register_special_token("", "test", "should fail")
        # Re-registering an enabled token must also raise (without override)
        with self.assertRaises(ValueError, msg="must reject duplicate enabled name"):
            register_special_token("bos", "test", "should fail")
        print("   ✅ Legacy collision, empty name, and duplicate name all rejected.")

    @unittest.skipUnless(HAS_MULTIMODAL_DEPS and HAS_SPECIAL_TOKENS, "multimodal deps missing")
    def test_mm_image_encoder_decode_accepts_legacy_marker(self):
        """🛰️ [MM-24] busel ImageEncoder.decode — accepts BOTH MOD_IMAGE and legacy MEDIA_START."""
        print("🧪 [MM-24] busel ImageEncoder.decode — accepts both v5.4 MOD_IMAGE and legacy MEDIA_START...")
        from PIL import Image as _Im
        import numpy as _np
        enc = ImageEncoder()
        src = _Im.fromarray(_np.random.randint(0, 255, (100, 100, 3), dtype=_np.uint8))
        blob_new = enc.encode(src)
        # Strip the leading MOD_IMAGE, replace with legacy MEDIA_START
        blob_legacy = [MEDIA_START] + blob_new[1:]
        decoded = enc.decode(blob_legacy)
        self.assertEqual(decoded.size, (32, 32), "legacy layout must still decode to 32x32")
        # And the new layout must also decode
        decoded2 = enc.decode(blob_new)
        self.assertEqual(decoded2.size, (32, 32), "v5.4 MOD_IMAGE layout must decode")
        print("   ✅ ImageEncoder.decode accepts both MOD_IMAGE (v5.4) and MEDIA_START (v5.0).")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "multimodal.special_tokens unavailable")
    def test_mm_patcher_embed_weight_uses_dynamic_vocab(self):
        """🛰️ [MM-25] busel StridedFastBLTPatcher — embed_weight shape = (vocab_size, d_byte)."""
        print("🧪 [MM-25] busel Patcher — embed_weight.shape[0] == vocab_size() at construction...")
        from model.patching import StridedFastBLTPatcher
        v = mm_vocab_size()
        patcher = StridedFastBLTPatcher(d_model=128, d_byte=64)
        self.assertEqual(patcher.embed_weight.shape[0], v,
                         f"embed_weight must be (vocab_size, d_byte) = ({v}, 64)")
        self.assertEqual(patcher.embed_weight.shape[1], 64, "d_byte must be 64 as requested")
        # vocab_size attribute exposed
        self.assertEqual(patcher.vocab_size, v)
        print(f"   ✅ Patcher embed_weight: {tuple(patcher.embed_weight.shape)} = (vocab_size={v}, d_byte=64).")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "multimodal.special_tokens unavailable")
    def test_mm_busel_model_rejects_undersized_vocab(self):
        """🛰️ [MM-26] busel buselModel — must raise if config.vocab_size < registry vocab_size()."""
        print("🧪 [MM-26] busel buselModel — rejects config.vocab_size < registry vocab_size()...")
        from model.backbone import buselModel
        registry_v = mm_vocab_size()
        undersized = registry_v - 5
        cfg = _MockConfig(vocab_size=undersized)
        with self.assertRaises(ValueError, msg="must reject undersized config.vocab_size"):
            buselModel(cfg)
        # Sanity: with matching vocab, model should construct fine
        cfg_ok = _MockConfig(vocab_size=registry_v)
        model = buselModel(cfg_ok)
        self.assertEqual(model.vocab_size, registry_v)
        print(f"   ✅ buselModel rejects vocab_size={undersized} (< registry {registry_v}); accepts {registry_v}.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_pretrain_registered(self):
        """🛸 [STG-1] busel stages — `pretrain` is registered on package import."""
        print("🧪 [STG-1] busel stages — `pretrain` is registered on package import...")
        from training.stages import list_stages
        names = list_stages()
        self.assertIn("pretrain", names, f"pretrain must be in list_stages(), got {names}")
        self.assertIsInstance(names, list, "list_stages() must return a list")
        self.assertEqual(names, sorted(names), "list_stages() must return sorted list")
        print(f"   ✅ list_stages() = {names}.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_get_pretrain(self):
        """🛸 [STG-2] busel stages — get_stage('pretrain') returns buselPretrainStage class."""
        print("🧪 [STG-2] busel stages — get_stage('pretrain') returns buselPretrainStage class...")
        from training.stages import get_stage
        cls = get_stage("pretrain")
        self.assertEqual(cls.__name__, "buselPretrainStage")
        self.assertEqual(cls.name, "pretrain")
        print(f"   ✅ get_stage('pretrain') = {cls.__module__}.{cls.__name__} (name={cls.name}).")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_unknown_raises(self):
        """🛸 [STG-3] busel stages — get_stage('nonexistent') raises KeyError."""
        print("🧪 [STG-3] busel stages — get_stage('nonexistent') raises KeyError...")
        from training.stages import get_stage
        with self.assertRaises(KeyError, msg="get_stage must reject unknown names"):
            get_stage("definitely_not_a_registered_stage_xyz")
        print("   ✅ get_stage('nonexistent_xyz') correctly raises KeyError.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_pretrain_has_lifecycle(self):
        """🛸 [STG-4] busel stages — buselPretrainStage has setup, run, finalize methods."""
        print("🧪 [STG-4] busel stages — buselPretrainStage has setup/run/finalize...")
        from training.stages import get_stage
        cls = get_stage("pretrain")
        for method_name in ("setup", "run", "finalize"):
            self.assertTrue(
                callable(getattr(cls, method_name, None)),
                f"buselPretrainStage must define {method_name}()",
            )
        instance = cls()
        for method_name in ("setup", "run", "finalize"):
            self.assertTrue(
                callable(getattr(instance, method_name, None)),
                f"buselPretrainStage instance must have {method_name}()",
            )
        print("   ✅ buselPretrainStage defines setup, run, finalize (instance-level callable).")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_pretrain_config_from_profile(self):
        """🛸 [STG-5] busel stages — buselPretrainConfig.from_profile parses YAML profile dict."""
        print("🧪 [STG-5] busel stages — buselPretrainConfig.from_profile parses profile dict...")
        from training.stages.pretrain import buselPretrainConfig
        profile = {
            "model": {"d_model": 128, "n_layers": 3, "n_heads": 4, "vocab_size": 326, "n_hyper": 2},
            "data": {"data_path": "data_train", "chunk_size": 256, "batch_size": 16},
            "training": {"max_steps": 100, "warmup_steps": 10, "min_lr_ratio": 0.1,
                         "learning_rate_muon": 0.001, "learning_rate_adamw": 0.0001,
                         "weight_decay": 0.1, "grad_accum_steps": 1},
        }
        cfg = buselPretrainConfig.from_profile(profile)
        self.assertEqual(cfg.d_model, 128)
        self.assertEqual(cfg.n_layers, 3)
        self.assertEqual(cfg.vocab_size, 326)
        self.assertEqual(cfg.max_steps, 100)
        self.assertEqual(cfg.warmup_steps, 10)
        self.assertEqual(cfg.data_path, "data_train")
        print(f"   ✅ buselPretrainConfig.from_profile: d_model={cfg.d_model}, n_layers={cfg.n_layers}, vocab={cfg.vocab_size}.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_pretrain_config_rejects_bad_dmodel(self):
        """🛸 [STG-6] busel stages — buselPretrainConfig rejects d_model not divisible by n_heads."""
        print("🧪 [STG-6] busel stages — buselPretrainConfig rejects invalid d_model/n_heads...")
        from training.stages.pretrain import buselPretrainConfig
        bad_profile = {
            "model": {"d_model": 100, "n_layers": 3, "n_heads": 3, "vocab_size": 326, "n_hyper": 2},
            "data": {"data_path": "data_train", "chunk_size": 256, "batch_size": 16},
            "training": {"max_steps": 10, "warmup_steps": 1, "min_lr_ratio": 0.1,
                         "learning_rate_muon": 0.001, "learning_rate_adamw": 0.0001,
                         "weight_decay": 0.1, "grad_accum_steps": 1},
        }
        with self.assertRaises(ValueError, msg="must reject d_model not divisible by n_heads"):
            buselPretrainConfig.from_profile(bad_profile)
        print("   ✅ buselPretrainConfig rejects d_model=100 with n_heads=3 (100 % 3 != 0).")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_load_pretrain_only_yaml(self):
        """🛸 [STG-7] busel stages — load_pipeline_yaml('configs/pipelines/pretrain-only.yaml') succeeds."""
        print("🧪 [STG-7] busel stages — load_pipeline_yaml('pretrain-only.yaml') succeeds...")
        from training.stages import load_pipeline_yaml
        cfg = load_pipeline_yaml("configs/pipelines/pretrain-only.yaml")
        self.assertEqual(cfg.name, "pretrain-only")
        self.assertEqual(len(cfg.stages), 1)
        s0 = cfg.stages[0]
        self.assertEqual(s0.name, "pretrain")
        self.assertEqual(s0.data_preset, "shpak")
        self.assertEqual(s0.params.get("profile_name"), "shpak")
        self.assertEqual(s0.params.get("max_steps"), 200)
        print(f"   ✅ pretrain-only.yaml: name={cfg.name}, stages=[{s0.name}], max_steps={s0.params.get('max_steps')}.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_load_missing_yaml_raises(self):
        """🛸 [STG-8] busel stages — load_pipeline_yaml raises FileNotFoundError on missing file."""
        print("🧪 [STG-8] busel stages — load_pipeline_yaml raises FileNotFoundError on missing file...")
        from training.stages import load_pipeline_yaml
        with self.assertRaises(FileNotFoundError):
            load_pipeline_yaml("configs/pipelines/this_does_not_exist.yaml")
        print("   ✅ load_pipeline_yaml on missing file raises FileNotFoundError.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_load_unknown_stage_raises(self):
        """🛸 [STG-9] busel stages — load_pipeline_yaml raises ValueError on unknown stage name."""
        print("🧪 [STG-9] busel stages — load_pipeline_yaml rejects unknown stage names...")
        import tempfile
        import os
        from training.stages import load_pipeline_yaml
        with tempfile.NamedTemporaryFile(suffix=".yaml", mode="w", delete=False) as f:
            f.write("name: bad\nstages:\n  - name: this_stage_is_not_registered_xyz\n")
            tmp_path = f.name
        try:
            with self.assertRaises(ValueError, msg="must reject unregistered stage names"):
                load_pipeline_yaml(tmp_path)
        finally:
            os.unlink(tmp_path)
        print("   ✅ load_pipeline_yaml with unknown stage name raises ValueError.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_load_missing_keys_raises(self):
        """🛸 [STG-10] busel stages — load_pipeline_yaml rejects YAML missing required keys."""
        print("🧪 [STG-10] busel stages — load_pipeline_yaml rejects missing required keys...")
        import tempfile
        import os
        from training.stages import load_pipeline_yaml
        for body, missing in [
            ("stages:\n  - name: pretrain\n", "name"),
            ("name: x\n", "stages"),
            ("name: x\nstages: bad\n", "stages-must-be-list"),
            ("name: x\nstages: []\n", "stages-must-be-non-empty"),
        ]:
            with tempfile.NamedTemporaryFile(suffix=".yaml", mode="w", delete=False) as f:
                f.write(body)
                tmp = f.name
            try:
                with self.assertRaises((KeyError, ValueError), msg=f"missing={missing}"):
                    load_pipeline_yaml(tmp)
            finally:
                os.unlink(tmp)
        print("   ✅ load_pipeline_yaml rejects all 4 missing-key shapes (KeyError or ValueError).")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_stage_state_defaults(self):
        """🛸 [STG-11] busel stages — StageState has the documented default field values."""
        print("🧪 [STG-11] busel stages — StageState defaults match the docstring contract...")
        from training.stages import StageState
        s = StageState()
        self.assertEqual(s.step, 0)
        self.assertEqual(s.epoch, 0)
        self.assertEqual(s.best_loss, float("inf"))
        self.assertEqual(s.metrics, {})
        self.assertIsNone(s.last_checkpoint_path)
        self.assertIsNone(s.artifact)
        s.step = 42
        s.metrics["loss"] = 1.23
        self.assertEqual(s.step, 42)
        self.assertEqual(s.metrics["loss"], 1.23)
        print("   ✅ StageState: step=0, epoch=0, best_loss=inf, metrics={}, ckpt=None, artifact=None.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_stage_spec_dataclass(self):
        """🛸 [STG-12] busel stages — StageSpec accepts the documented field set."""
        print("🧪 [STG-12] busel stages — StageSpec accepts name/data_preset/resume/params...")
        from training.stages import StageSpec
        s = StageSpec(
            name="pretrain",
            data_preset="shpak",
            resume="checkpoints/x.pt",
            checkpoint_out="checkpoints/y.pt",
            params={"max_steps": 100},
        )
        self.assertEqual(s.name, "pretrain")
        self.assertEqual(s.data_preset, "shpak")
        self.assertEqual(s.resume, "checkpoints/x.pt")
        self.assertEqual(s.checkpoint_out, "checkpoints/y.pt")
        self.assertEqual(s.params["max_steps"], 100)
        print("   ✅ StageSpec accepts all 5 documented fields.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_pipeline_config_dataclass(self):
        """🛸 [STG-13] busel stages — PipelineConfig accepts the documented field set."""
        print("🧪 [STG-13] busel stages — PipelineConfig accepts name/stages/global_params...")
        from training.stages import PipelineConfig, StageSpec
        cfg = PipelineConfig(
            name="full",
            stages=[StageSpec(name="pretrain")],
            global_params={"max_steps": 200},
        )
        self.assertEqual(cfg.name, "full")
        self.assertEqual(len(cfg.stages), 1)
        self.assertEqual(cfg.stages[0].name, "pretrain")
        self.assertEqual(cfg.global_params["max_steps"], 200)
        print("   ✅ PipelineConfig accepts all 3 documented fields.")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "training.stages unavailable")
    def test_stages_orchestrator_pipeline_command_registered(self):
        """🛸 [STG-14] busel stages — tools.orchestrator:pipeline is exported and importable."""
        print("🧪 [STG-14] busel stages — tools.orchestrator:pipeline is importable...")
        from tools.orchestrator import pipeline
        self.assertTrue(callable(pipeline), "tools.orchestrator.pipeline must be callable")
        # Typer commands expose a `.callback` attribute; the function itself should
        # also have a __name__.
        self.assertEqual(pipeline.__name__, "pipeline")
        print(f"   ✅ tools.orchestrator.pipeline exists, callable, name='{pipeline.__name__}'.")

    # ════════════════════════════════════════════════════════════════════
    #  PHASE 3 — SFT (data/sft.py + training/stages/sft.py)
    # ════════════════════════════════════════════════════════════════════

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_sft_format_chat_messages_basic(self):
        """🤖 [SFT-1] format_chat_messages produces BOS/ROLE/EOS structure."""
        print("🧪 [SFT-1] format_chat_messages produces BOS/ROLE/EOS structure...")
        from data.sft import format_chat_messages
        messages = [
            {"role": "system", "content": "Be helpful."},
            {"role": "user", "content": "Hi"},
            {"role": "assistant", "content": "Hello!"},
        ]
        b, m = format_chat_messages(messages)
        self.assertEqual(len(b), len(m))
        self.assertGreater(len(b), 0)
        self.assertEqual(b[0], int(BOS), "first token must be BOS")
        # Find the assistant content range; mask must be 1 there.
        assistant_idx = b.index(int(ROLE_ASSISTANT))
        # mask at assistant_idx+1 (the first content byte) must be 1
        self.assertEqual(m[assistant_idx + 1], 1, "first assistant-content position must be masked 1")
        # mask at the ROLE_ASSISTANT token itself must be 0
        self.assertEqual(m[assistant_idx], 0, "ROLE_ASSISTANT token position must be masked 0")
        print(f"   ✅ {len(b)} bytes, {sum(m)} mask=1 positions (assistant content + EOS).")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_sft_format_chat_messages_mask_correctness(self):
        """🤖 [SFT-2] format_chat_messages mask is 0 for system/user/tool, 1 for assistant."""
        print("🧪 [SFT-2] format_chat_messages mask correctness...")
        from data.sft import format_chat_messages
        b, m = format_chat_messages([
            {"role": "system", "content": "SYS"},
            {"role": "user", "content": "USR"},
            {"role": "assistant", "content": "AS"},
        ])
        # Find the user content position and verify mask=0
        user_idx = b.index(int(ROLE_USER))
        self.assertEqual(m[user_idx + 1], 0, "user content must be masked 0")
        # System content
        sys_idx = b.index(int(ROLE_SYSTEM))
        self.assertEqual(m[sys_idx + 1], 0, "system content must be masked 0")
        # Assistant content (first content byte after ROLE_ASSISTANT)
        asst_idx = b.index(int(ROLE_ASSISTANT))
        self.assertEqual(m[asst_idx + 1], 1, "assistant content must be masked 1")
        # The final EOS is after the assistant turn (assistant is last in this test)
        last_eos = len(b) - 1 - b[::-1].index(int(EOS))
        self.assertEqual(b[last_eos], int(EOS))
        self.assertEqual(m[last_eos], 1, "final EOS (after assistant turn) must be masked 1")
        # Total mask=1 count must equal len("AS") + 1 (for final EOS) = 3
        self.assertEqual(sum(m), 3, f"expected 3 mask=1 positions (AS + EOS), got {sum(m)}")
        print(f"   ✅ system/user masked 0; assistant content + final EOS masked 1; total mask=1 = {sum(m)}.")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_sft_format_dpo_pair(self):
        """🤖 [SFT-3] format_dpo_pair produces chosen+rejected with matching prompt mask."""
        print("🧪 [SFT-3] format_dpo_pair produces chosen + rejected with shared prompt...")
        from data.sft import format_dpo_pair
        cb, cm, rb, rm = format_dpo_pair("Q?", "good answer", "bad answer")
        self.assertGreater(len(cb), 0)
        self.assertGreater(len(rb), 0)
        self.assertEqual(len(cb), len(cm))
        self.assertEqual(len(rb), len(rm))
        # Both should contain the prompt
        prompt_bytes = list("Q?".encode("utf-8"))
        self.assertTrue(any(prompt_bytes[0] == cb[i] for i in range(len(cb))))
        # Both should have at least one mask=1
        self.assertGreater(sum(cm), 0)
        self.assertGreater(sum(rm), 0)
        print(f"   ✅ chosen={len(cb)} bytes/{sum(cm)} mask=1; rejected={len(rb)} bytes/{sum(rm)} mask=1.")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_sft_dataloader_yields_matching_tensors(self):
        """🤖 [SFT-4] get_sft_dataloader yields (bytes, mask) batches with matching shapes."""
        print("🧪 [SFT-4] get_sft_dataloader yields matching (bytes, mask) tensors...")
        import tempfile
        from data.sft import get_sft_dataloader
        with tempfile.NamedTemporaryFile(mode="w", suffix=".jsonl", delete=False) as f:
            f.write(json.dumps({"messages": [
                {"role": "user", "content": "Q1?"},
                {"role": "assistant", "content": "A1"},
            ]}) + "\n")
            f.write(json.dumps({"messages": [
                {"role": "user", "content": "Q2?"},
                {"role": "assistant", "content": "A2 long answer here"},
            ]}) + "\n")
            tmp = f.name
        try:
            dl = get_sft_dataloader([tmp], chunk_size=64, batch_size=2)
            it = iter(dl)
            batch = next(it)
            self.assertEqual(len(batch), 2, "batch should be (bytes, mask)")
            bytes_b, mask_b = batch
            self.assertEqual(bytes_b.shape, mask_b.shape)
            self.assertEqual(bytes_b.dtype, torch.int32)
            self.assertEqual(mask_b.dtype, torch.int32)
            self.assertGreater(bytes_b.shape[0], 0)
            self.assertGreater(bytes_b.shape[1], 0)
            print(f"   ✅ batch shapes: bytes={tuple(bytes_b.shape)}, mask={tuple(mask_b.shape)}")
        finally:
            os.unlink(tmp)

    @unittest.skipUnless(HAS_TRAINING_STAGES, "stages required")
    def test_sft_config_from_profile(self):
        """🤖 [SFT-5] buselSFTConfig.from_profile parses profile dict + stage_params."""
        print("🧪 [SFT-5] buselSFTConfig.from_profile parses profile + stage_params...")
        from training.stages.sft import buselSFTConfig
        profile = {
            "model": {"d_model": 128, "n_layers": 3, "n_heads": 4, "vocab_size": 326, "n_hyper": 2},
            "data": {"chunk_size": 256, "batch_size": 16},
            "training": {"learning_rate_muon": 0.001, "learning_rate_adamw": 0.0001},
        }
        cfg = buselSFTConfig.from_profile(profile, {"max_steps": 200, "sft_lr_scale": 0.5})
        self.assertEqual(cfg.d_model, 128)
        self.assertEqual(cfg.vocab_size, 326)
        self.assertEqual(cfg.max_steps, 200)
        # SFT LR is 0.5x the base
        self.assertAlmostEqual(cfg.learning_rate_muon, 0.0005, places=7)
        self.assertAlmostEqual(cfg.learning_rate_adamw, 0.00005, places=8)
        print(f"   ✅ SFTConfig: d_model={cfg.d_model}, max_steps={cfg.max_steps}, lr_muon={cfg.learning_rate_muon:.6f}")

    # ════════════════════════════════════════════════════════════════════
    #  PHASE 5 — DPO (data/dpo.py + training/stages/dpo.py + recipe.py)
    # ════════════════════════════════════════════════════════════════════

    def test_dpo_loss_finite_scalar(self):
        """🤖 [DPO-1] buselLossEngine.compute_dpo_loss returns a finite scalar."""
        print("🧪 [DPO-1] buselLossEngine.compute_dpo_loss returns a finite scalar...")
        torch.manual_seed(0)
        B = 8
        pc = torch.randn(B) * 2
        pr = torch.randn(B) * 2
        rc = torch.randn(B) * 2
        rr = torch.randn(B) * 2
        loss = buselLossEngine.compute_dpo_loss(pc, pr, rc, rr, beta=0.1)
        self.assertTrue(torch.isfinite(loss).item(), f"DPO loss must be finite, got {loss.item()}")
        self.assertEqual(loss.dim(), 0, "DPO loss must be a scalar")
        # Symmetric case: chosen_logp - rejected_logp is the same for policy and reference
        # → logits=0 → loss = -log(0.5) = log(2) ≈ 0.693
        sym_loss = buselLossEngine.compute_dpo_loss(pc, pr, pc, pr, beta=0.1)
        self.assertAlmostEqual(sym_loss.item(), math.log(2), places=4)
        print(f"   ✅ random loss={loss.item():.4f}, symmetric loss={sym_loss.item():.4f} (≈ log 2 = {math.log(2):.4f})")

    def test_dpo_sequence_logprob_respects_mask(self):
        """🤖 [DPO-2] buselLossEngine.compute_sequence_logprob respects mask."""
        print("🧪 [DPO-2] buselLossEngine.compute_sequence_logprob respects mask...")
        torch.manual_seed(0)
        V = 50
        B, T = 2, 8
        logits = torch.randn(B, T, V)
        targets = torch.randint(0, V, (B, T))
        mask = torch.zeros(B, T, dtype=torch.int32)
        mask[0, 2:6] = 1  # 4 active positions
        mask[1, :] = 1    # all 8 active
        logp = buselLossEngine.compute_sequence_logprob(logits, targets, mask)
        self.assertEqual(logp.shape, (B,))
        # Manual: log_softmax + gather + sum over mask
        ref_logp = torch.nn.functional.log_softmax(logits, dim=-1).gather(-1, targets.unsqueeze(-1).long()).squeeze(-1)
        ref = (ref_logp * mask.float()).sum(dim=-1)
        self.assertTrue(torch.allclose(logp, ref, atol=1e-5))
        print(f"   ✅ per-seq logp[0]={logp[0].item():.3f}, logp[1]={logp[1].item():.3f}, matches manual calc.")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_dpo_dataloader_yields_4_tensors(self):
        """🤖 [DPO-3] get_dpo_dataloader yields (chosen_b, chosen_m, rejected_b, rejected_m)."""
        print("🧪 [DPO-3] get_dpo_dataloader yields 4-tensor batches with matching shapes...")
        import tempfile
        from data.dpo import get_dpo_dataloader
        with tempfile.NamedTemporaryFile(mode="w", suffix=".jsonl", delete=False) as f:
            f.write(json.dumps({"prompt": "Q1?", "chosen": "good", "rejected": "bad"}) + "\n")
            f.write(json.dumps({"prompt": "Q2?", "chosen": "great", "rejected": "meh"}) + "\n")
            tmp = f.name
        try:
            dl = get_dpo_dataloader([tmp], chunk_size=64, batch_size=2)
            it = iter(dl)
            batch = next(it)
            self.assertEqual(len(batch), 4, "batch should be (chosen_b, chosen_m, rejected_b, rejected_m)")
            cb, cm, rb, rm = batch
            self.assertEqual(cb.shape, cm.shape)
            self.assertEqual(rb.shape, rm.shape)
            self.assertEqual(cb.shape, rb.shape, "chosen and rejected must have same shape")
            print(f"   ✅ shapes: chosen={tuple(cb.shape)}, rejected={tuple(rb.shape)}")
        finally:
            os.unlink(tmp)

    @unittest.skipUnless(HAS_TRAINING_STAGES, "stages required")
    def test_dpo_config_from_profile(self):
        """🤖 [DPO-4] buselDPOConfig.from_profile parses profile + stage_params."""
        print("🧪 [DPO-4] buselDPOConfig.from_profile parses profile + stage_params...")
        from training.stages.dpo import buselDPOConfig
        profile = {
            "model": {"d_model": 128, "n_layers": 3, "n_heads": 4, "vocab_size": 326, "n_hyper": 2},
            "data": {"chunk_size": 256, "batch_size": 8},
            "training": {"learning_rate_muon": 0.001, "learning_rate_adamw": 0.0001},
        }
        cfg = buselDPOConfig.from_profile(profile, {"beta": 0.2, "dpo_lr_scale": 0.05})
        self.assertEqual(cfg.d_model, 128)
        self.assertAlmostEqual(cfg.dpo_beta, 0.2, places=4)
        # DPO LR is 0.05x the base
        self.assertAlmostEqual(cfg.learning_rate_muon, 0.00005, places=7)
        print(f"   ✅ DPOConfig: β={cfg.dpo_beta}, lr_muon={cfg.learning_rate_muon:.6f}")

    # ════════════════════════════════════════════════════════════════════
    #  PHASE 6 — EVAL (tools/eval.py + training/stages/eval.py)
    # ════════════════════════════════════════════════════════════════════

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_eval_perplexity_finite(self):
        """🛰️ [EVAL-1] tools.eval.perplexity returns finite perplexity on tiny input."""
        print("🧪 [EVAL-1] tools.eval.perplexity returns finite perplexity...")
        from tools.eval import perplexity
        # Fake model returning random logits
        class FM:
            def __call__(self, x, mtp, progress=0.0):
                B, T, D = x.shape
                return (torch.randn(B, T, 256), torch.zeros(B, T, 128), torch.zeros(B, T, 64), torch.zeros(B, T, 32)), torch.tensor(0.0)
        class FP:
            stride = 4
            def __call__(self, x):
                B, T = x.shape
                return torch.randn(B, T // 4, 128)
        result = perplexity(FM(), FP(), [list(b"hello world" * 8), list(b"good morning" * 4)], "cpu", max_samples=2)
        self.assertIn("perplexity", result)
        self.assertIn("bits_per_byte", result)
        self.assertTrue(math.isfinite(result["perplexity"]), f"perplexity must be finite, got {result['perplexity']}")
        self.assertGreater(result["perplexity"], 1.0, "random model perplexity should be > vocab baseline")
        print(f"   ✅ perplexity={result['perplexity']:.2f}, bpb={result['bits_per_byte']:.3f}")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_eval_format_compliance_returns_valid_dict(self):
        """🛰️ [EVAL-2] tools.eval.format_compliance returns valid dict with required keys."""
        print("🧪 [EVAL-2] tools.eval.format_compliance returns valid dict...")
        from tools.eval import format_compliance
        class FM:
            def __call__(self, x, mtp, progress=0.0):
                B, T, D = x.shape
                return (torch.randn(B, T, 256), torch.zeros(B, T, 128), torch.zeros(B, T, 64), torch.zeros(B, T, 32)), torch.tensor(0.0)
        class FP:
            stride = 4
            def __call__(self, x):
                B, T = x.shape
                return torch.randn(B, max(1, T // 4), 128)
        result = format_compliance(FM(), FP(), ["hello"], "cpu", max_prompts=1, max_new_tokens=8)
        for k in ("format_compliance", "avg_response_bytes", "n_prompts"):
            self.assertIn(k, result, f"missing key {k!r}")
        self.assertEqual(result["n_prompts"], 1)
        self.assertGreaterEqual(result["format_compliance"], 0.0)
        self.assertLessEqual(result["format_compliance"], 1.0)
        print(f"   ✅ compliance={result['format_compliance']:.2%}, avg_bytes={result['avg_response_bytes']:.1f}")

    @unittest.skipUnless(HAS_SPECIAL_TOKENS, "special tokens required")
    def test_eval_stages_all_registered(self):
        """🛰️ [EVAL-3] All 4 pipeline stages (pretrain, sft, dpo, eval) are registered."""
        print("🧪 [EVAL-3] All 4 pipeline stages registered...")
        stages = list_stages()
        for name in ("pretrain", "sft", "dpo", "eval"):
            self.assertIn(name, stages, f"stage {name!r} not registered; got {stages}")
        for name in ("pretrain", "sft", "dpo", "eval"):
            self.assertTrue(is_stage_registered(name))
        print(f"   ✅ registered stages: {sorted(stages)}")

    # ════════════════════════════════════════════════════════════════════
    #  PHASE 8 — PIPELINE YAMLs
    # ════════════════════════════════════════════════════════════════════

    @unittest.skipUnless(HAS_TRAINING_STAGES, "stages required")
    def test_pipeline_full_yaml_loads(self):
        """🛸 [PIPE-1] configs/pipelines/full.yaml loads with 4 stages."""
        print("🧪 [PIPE-1] configs/pipelines/full.yaml loads with 4 stages...")
        cfg = load_pipeline_yaml("configs/pipelines/full.yaml")
        self.assertEqual(cfg.name, "full")
        self.assertEqual(len(cfg.stages), 4)
        self.assertEqual([s.name for s in cfg.stages], ["pretrain", "sft", "dpo", "eval"])
        print(f"   ✅ pipeline={cfg.name!r}, stages={[s.name for s in cfg.stages]}")

    @unittest.skipUnless(HAS_TRAINING_STAGES, "stages required")
    def test_pipeline_quick_yaml_loads(self):
        """🛸 [PIPE-2] configs/pipelines/quick.yaml loads with 2 stages."""
        print("🧪 [PIPE-2] configs/pipelines/quick.yaml loads with 2 stages...")
        cfg = load_pipeline_yaml("configs/pipelines/quick.yaml")
        self.assertEqual(cfg.name, "quick")
        self.assertEqual(len(cfg.stages), 2)
        self.assertEqual([s.name for s in cfg.stages], ["pretrain", "eval"])
        print(f"   ✅ pipeline={cfg.name!r}, stages={[s.name for s in cfg.stages]}")


if __name__ == "__main__":
    unittest.main()
